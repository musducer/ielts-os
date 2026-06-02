import re
from docx import Document

def compile_markdown_to_docx(md_filepath, output_filename="IELTS_Converted.docx"):
    with open(md_filepath, "r", encoding="utf-8") as f:
        content = f.read()

    # ==========================================
    # BỘ LỌC TẨY RỬA RÁC TRÍCH DẪN (AI CITATIONS)
    # ==========================================
    # 1. Xóa toàn bộ các tag bắt đầu trích dẫn: 
    clean_content = re.sub(r'\[cite_start\]\s*', '', content)
    
    # 2. Xóa toàn bộ các tag đánh số nguồn:,, v.v.
    clean_content = re.sub(r'\s*\]*\]', '', clean_content)

    doc = Document()
    for line in clean_content.split("\n"):
        text = line.strip()
        if not text:
            doc.add_paragraph("")
            continue
        
        p = doc.add_paragraph()
        
        # Xử lý In đậm cơ bản (**text**)
        parts = re.split(r"(\*\*.*?\*\*)", text)
        for part in parts:
            if part.startswith("**") and part.endswith("**"):
                p.add_run(part[2:-2]).bold = True
            else:
                p.add_run(part)

    doc.save(output_filename)
    print(f"✅ Đã đúc thành công: {output_filename}. Bản Word đã sạch sẽ và sẵn sàng nạp vào IELTS-OS!")

if __name__ == "__main__":
    compile_markdown_to_docx("data.md", "V5T6_Passage1.docx")