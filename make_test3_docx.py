# -*- coding: utf-8 -*-
"""Generate Test 3.docx theo format backend parse"""
from docx import Document

doc = Document()

def p(text="", bold=False):
    para = doc.add_paragraph()
    run = para.add_run(text)
    if bold:
        run.bold = True
    return para

# META
p("[TITLE] Test 3 - IELTS Academic Reading")
p("[TIME] 60")
p("[TYPE] Reading")

# ============================================================
# PASSAGE 1
# ============================================================
p("[PASSAGE]")
p("The development of plastics", bold=True)
p("The first plastics were developed as a substitute for natural rubber. Chemically, rubber is a polymer-a compound containing large molecules that are formed by the bonding of many smaller, simpler units, repeated over and over again. The same bonding principle-polymerization-is the basis of the creation of a huge range of plastics by the chemical industry.")
p("The first plastic was developed as a result of a competition in the USA. In the 1860s, $10,000 was offered to anybody who could replace ivory-supplies of which were declining-with something equally good as a material for making billiard balls. The prize was won by John Wesley Hyatt, with a material called celluloid. Celluloid was made by dissolving cellulose, a carbohydrate obtained from plants, in a solution of camphor dissolved in ethanol. This new material rapidly found other applications in the manufacture of everyday products such as knife handles and detachable collars and cuffs. But perhaps the best-known celluloid product was photographic film, without which the film industry could never have taken off at the end of the 19th century.")
p("Celluloid can be repeatedly softened and reshaped by heat, and is known as a thermoplastic. In 1907, Leo Baekeland (1863-1944), a Belgian chemist working in the USA, invented a different kind of plastic, by causing phenol and formaldehyde to react together. Baekeland called it Bakelite, and it was the first of the thermosets-plastic that can be cast and moulded while hot, but cannot be softened by heat and reshaped once they have set. Bakelite was a good insulator, and was resistant to water and acid. With these properties it was soon being used in the manufacture of electrical switches as well as a variety of domestic items.")
p("As the century went on, the range of newly developed plastic increased. Chemists began looking for other small molecules that could be strung together to make polymers. In the 1930s, chemists in Britain discovered that the gas ethylene would polymerize under heat and pressure to form a thermoplastic they called polythene. Polypropylene followed in the 1950s. Both are used to make bottles, pipes and plastic bags. A small change in the starting material-replacing a hydrogen atom in ethylene with a chlorine atom-produced rigid PVC (polyvinyl chloride), a fireproof plastic suitable for drains and gutters. By adding certain chemicals, a soft form of PVC can be produced, suitable as a substitute for rubber in items such as waterproof clothing. A closely related plastic is Teflon or PTFE (polytetrafluoroethylene). It produces very little friction, making it ideal for products such as non-stick frying pans.")
p("Polystyrene, a hard, clear material like glass, was developed during the 1930s in Germany, and its applications included food containers and toys. Expanded polystyrene is rigid and is widely used in packaging and insulation. Polyurethane, developed in the same country, was commonly produced as a foam which was very useful in the production of insulating materials.")
p("In the 1930s, the first of the man-made fibres was created-nylon. Its inventor was a chemist called Wallace Carothers (1896-1937), who worked for the Du Pont company in the USA. He found that under the right conditions two particular chemicals would form a polymer that could be pumped out through holes and then stretched to form long glossy threads that could be woven like silk. Its first use was to make parachutes for the US armed forces in World War II. In the postwar years, it completely replaced silk in the manufacture of stockings.")
p("Many other synthetic fibres joined nylon, including Orion, Acrilan, and Terylene. Today most garments are made of a blend of natural fibres, such as cotton and wool, and man-made fibres that make fabrics easier to look after.")
p("Despite its enormous usefulness, plastic has its drawbacks. In fact one of its great strengths-its indestructibility-is its greatest disadvantage. Beaches all over the world, even on the remotest island, are littered with plastic bottles that nothing can destroy. Nor is it very easy to recycle plastics, as different types of plastic are often found in the same items and call for different treatments.")
p("Plastics can be made biodegradable by incorporating into their structure a material such as starch, which is attacked by bacteria and causes the plastic to fall apart. Other materials can be incorporated that gradually decay in sunlight-although bottles made of such materials have to be stored in the dark, to ensure they do not disintegrate before they have been used.")

# QUESTIONS - Passage 1
p("[QUESTIONS]")

# Q1-7: Table completion -> SHORT_ANSWER -> BLANK, bang nguyen trong [CONTEXT], gap [1]..[7]
p("[SHORT_ANSWER]")
p("Questions 1-7")
p("Complete the table below.")
p("Write NO MORE THAN TWO WORDS from the passage for each answer.")
p("Early types of plastic", bold=True)

p("[CONTEXT]")
TBL = [
    ["Name", "Date", "Country of origin", "Properties", "Common uses"],
    ["Celluloid", "1860s", "USA", "Can be softened and reshaped by heat", "billiard balls (original use); cutlery; clothing; [1]"],
    ["[2]", "1907", "USA", "Can't be softened after setting; good insulator; resistant to water and acid", "[3]; Household object"],
    ["Polythene", "1930s", "[4]", "Can be softened and reshaped by heat", "bottles; pipes; plastic bags"],
    ["Polypropylene", "1950s", "", "", "bottles; pipes; plastic bags"],
    ["Rigid PVC", "", "", "Is [5]", "external piping"],
    ["Soft PVC", "", "", "", "outdoor clothing"],
    ["Polystyrene", "1930s", "Germany", "Resembles [6]", "food containers; toy"],
    ["Polyurethane", "", "Germany", "Usually manufactured as a [7]", "insulation"],
]
_t = doc.add_table(rows=len(TBL), cols=5)
_t.style = "Table Grid"
for _ri, _row in enumerate(TBL):
    for _ci, _val in enumerate(_row):
        _t.cell(_ri, _ci).text = _val
p("[/CONTEXT]")

p("1.")
p("*photographic film")
p("2.")
p("*Bakelite")
p("3.")
p("*electrical switches")
p("4.")
p("*Britain")
p("5.")
p("*fireproof")
p("6.")
p("*glass")
p("7.")
p("*foam")

# Q8-13: TRUE/FALSE/NOT GIVEN -> CHOICE
p("[CHOICE]")
p("Questions 8-13")
p("Choose TRUE if the statement agrees with the information given in the text, choose FALSE if the statement contradicts the information, or choose NOT GIVEN if there is no information on this. True / False / Not Given")

p("8. The chemical structure of rubber is very different to that of plastics.")
p("*FALSE")

p("9. John Wesley Hyatt was an industrial chemist.")
p("*NOT GIVEN")

p("10. Celluloid and Bakelite react in the same way to heat.")
p("*FALSE")

p("11. If an object is made of several plastics, these prove hard to break down and reuse.")
p("*TRUE")

p("12. Adding starch to plastic makes it more durable.")
p("*FALSE")

p("13. Containers which are designed to decompose need particular storage conditions.")
p("*TRUE")

# ============================================================
# PASSAGE 2
# ============================================================
p("[PASSAGE]")
p("Lean Production innovation - in manufacturing systems", bold=True)
p("[HEADING_SLOT]")
p("After the First World War, car makers Henry Ford and Arthur Sloan of General Motors moved world manufacturing from centuries of craft production into the age of mass production. Largely as a result of this, the United States soon dominated the world economy. After the Second World War, and approximately a hundred years after Japan opened up to the modern world, Eiji Toyoda and Taiichi Ohno pioneered the concept of lean production at the Toyota car company. And now, although superimposing the method on existing mass production systems causes pain and upheaval, manufacturers around the world are trying to embrace this innovative system.")
p("[HEADING_SLOT]")
p("Perhaps the best way to describe lean production is to compare it with the two other major manufacturing systems: craft production and mass production. The craft producer uses highly skilled workers and simple but flexible tools to make exactly what the customer asks for - one item at a time. A present day example of this method is the customised production of a few exotic sports cars. The concept of craft production remains very popular, but the problem with it is obvious. Goods produced by the craft method - as cars once exclusively were - cost too much for most of us to afford. So at the beginning of the twentieth century, mass production was developed as an alternative method. The mass producer uses narrowly skilled professionals to design products which are then made by unskilled or semi-skilled workers, using expensive, single-purpose machines. These churn out standardised products in very high volumes. Because the machinery costs so much, and is so intolerant of disruption, the mass producer keeps standard designs in production for as long as possible. The result is that the customer gets lower costs, but at the expense of variety, and by means of work methods which most employees find boring and dispiriting. By contrast, the lean production system combines the advantages of craft and mass production, while avoiding the high cost of the former and the rigidity of the latter. Towards this end, companies appoint teams of multi-skilled workers to all levels of the organisation, and use highly flexible and increasingly automated machines to produce goods in enormous volume and variety.")
p("[HEADING_SLOT]")
p("Lean production is so called because, compared with mass production, it uses less of everything - half the human effort in the factory, half the manufacturing space, half the investment in tools, and half the engineering hours to develop the new product. It also results in far fewer defects.")
p("[HEADING_SLOT]")
p("Perhaps the most striking contrast between mass and lean production systems lies in their production standards. Mass producers set a limited goal for themselves: 'good enough', which translates into an acceptable number of defects, a maximum acceptable number of inventories, and a narrow range of standardised products. Lean producers, on the other hand, are unwilling to compromise standards in any of these areas.")
p("[HEADING_SLOT]")
p("Although cost reduction is the primary objective of the lean production system, it must meet three other intermediate objectives in order to achieve this: quantity control, quality assurance and respect for humanity. Firstly, the system must be able to adapt to daily and monthly fluctuations in demand. Secondly, each separate process must supply only good units to the subsequent process. Thirdly, in as far as the system uses human resources to attain its cost objectives, respect for human needs must be cultivated. It should be emphasised that none of these three objectives can be achieved separately.")
p("[HEADING_SLOT]")
p("The continuous flow of lean manufacturing production relies on two practical mechanisms: just-in-time and autonamation. Just-in-time means, for example, that in the process of assembling the parts to build a car, components from the preceding process should arrive at the next part of the line at exactly the right time and in the correct quantities. If just-in-time is fully realised throughout the company, then superfluous inventories are completely eliminated from the factory, making stores or warehouses unnecessary. However, relying solely on a central planning approach to control schedules for all stages of the production process simultaneously is very difficult in the case of cars, which consist of thousands of parts. So the lean system looks at the production flow in reverse; in other words, employees go to the preceding process to withdraw the necessary quantity of units at the appropriate time. The preceding process must produce only sufficient quantities of units to replace those that have been withdrawn, and in turn withdraws the requisite number of components from the process that precedes it.")
p("[HEADING_SLOT]")
p("Autonamation is the automatic checking for abnormalities in the production process. In order to realise just-in-time perfectly, only units which are in perfect condition must be allowed to flow to the next process, and this flow must be regular and uninterrupted. In other words, quality control must coexist with just-in-time procedures throughout the system. Autonamation involves building in a mechanism to prevent the multiplication of defects in machines or product lines. For example, in Toyota factories almost all the machines have been fitted with stopping devices, and the concept of autonamation has been extended to manual production lines. If something abnormal happens there, the worker pushes a button to stop the whole line, and lights, which are hung so high in the factory that they are visible to everyone, indicate the position of the problem.")

# QUESTIONS - Passage 2
p("[QUESTIONS]")

# Q14-20: Matching Headings -> MATCHING -> DRAG_DROP_HEADING
p("[MATCHING]")
p("Questions 14-20")
p("Reading Passage 2 has seven paragraphs.")
p("Choose the correct heading for each paragraph from the list of headings below.")
p("List of Headings")
p("i. Global resistance to lean manufacturing")
p("ii. The historical context")
p("iii. Procedures for controlling quality")
p("iv. The pros and cons of different production systems")
p("v. The impact on profits")
p("vi. Procedures for controlling supply")
p("vii. The origin of the term")
p("viii. A crucial difference in levels of quality")
p("ix. Working conditions")
p("x. Interdependent strategies for controlling expenditure")

p("14. Paragraph 14")
p("*ii")

p("15. Paragraph 15")
p("*iv")

p("16. Paragraph 16")
p("*vii")

p("17. Paragraph 17")
p("*viii")

p("18. Paragraph 18")
p("*x")

p("19. Paragraph 19")
p("*vi")

p("20. Paragraph 20")
p("*iii")

# Q21-26: Sentence completion -> SHORT_ANSWER -> BLANK
p("[SHORT_ANSWER]")
p("Questions 21-26")
p("Complete the sentences below.")
p("Write NO MORE THAN THREE WORDS from the passage for each answer.")

p("21. A small number of unusual cars are still produced by the ________ method.")
p("*craft")

p("22. Lean production requires staff who are ________.")
p("*multi-skilled")

p("23. Lean production employs fewer people, and uses less ________, equipment and time.")
p("*manufacturing space")

p("24. Storage facilities are not needed if a procedure known as ________ is implemented in the lean production method.")
p("*just-in-time")

p("25. Autonamation is a procedure for spotting any ________ in the products on a production line.")
p("*abnormalities")

p("26. At Toyota factories, ________ are suspended above manual production lines in order to show where production has to be halted.")
p("*lights")

# ============================================================
# PASSAGE 3
# ============================================================
p("[PASSAGE]")
p("How Does Watching Sport Influence the Brain?", bold=True)
p("A During the epoch when the poet Homer unveiled his grand epic, the ancient Greeks initiated a festival where men contended in a race of roughly 200 metres. The victor would be rewarded with a symbolic olive branch. This event, known as the Olympics, has transcended its humble origins to embody the human quest for progress. The captivating nature of these games raises an intriguing question: what exactly draws us in as spectators? Is it the thrill of competition, the beauty of human movement, or something happening within our brains?")
p("B In 1996, three Italian neuroscientists-Giacomo Rizzolatti, Leonardo Fogassi, and Vittorio Gallese conducted research on the premotor cortex of monkeys and made a groundbreaking discovery. They revealed a cluster of cells that act as a repository for motor actions, similar to 'grammar of movement'. These cellular networks represent the 'sentences' of full-body motions that our brains meticulously retain and refine. Take the art of swinging a golf club as an example. For those who only watch the Masters Tournament on television, playing golf may appear deceptively effortless. However, for novices, smoothly manoeuvring the asymmetrical metal club presents a nearly impossible challenge. This is because novices rely on conscious effort, utilising brain regions adjacent to the premotor cortex. Experts, by contrast, instinctively execute perfectly balanced movements, as their actions are deeply ingrained and intertwined with neurons in the premotor cortex. Their swing occurs with the serenity of a refined autopilot, requiring no conscious control. This neural distinction between novices and experts provides insight into how long-term training can physically reshape the brain.")
p("C These neurons in the premotor cortex not only help explain why some athletes seem to have exceptional skills, but they also exhibit a wondrous characteristic that prompted Rizzolatti, Fogassi, and Gallese to bestow upon them the exalted label of 'mirror neuron'. They point out that mirror neurons are activated when a primate performs a particular action, such as grasping or holding an object, or when it witnesses another individual performing a comparable action. Humans possess an even more intricate mirror neuron system that reflects the external world in our brains, enabling us to internalise the behaviours of others. Nonetheless, these cells necessitate what scientists refer to as 'goal-orientated movements' to activate. If we stare at a photograph, a static image of a runner striding ahead, our mirror neurons remain utterly silent. It is solely when the runner is in action, be it in running, moving, or sprinting, that they are ignited. Movement, not stillness, is what speaks to our brains.")
p("D Electrophysiological studies have shown that when we observe a golfer or a runner in action, our own premotor cortex mirror neurons respond as if we were participating in the activity. This neural mirroring was initially identified in 1954 by French physiologists Gastaut and Bert, who noticed specific alpha and mu brain wave patterns in humans. The mu signal plays a role in neural mirroring, remaining active during periods of bodily stillness but dissipating when we undertake active pursuits like playing a sport or changing TV channels. Notably, the mu signal also remains subdued while we observe someone else in action, such as on television, which alludes to the influence of mirror neurons. This suggests that our brains do not simply passively register what we see, but actively simulate it, engaging systems normally reserved for action.")
p("E Rizzolatti, Fogassi, and Gallese have coined the term 'direct matching hypothesis' to delineate the concept of mirror neurons. They asserted that to comprehend the movements of sports stars, we need to align the visual representation of the observed action with our own motor representation of the same action. According to this theory, watching an Olympic athlete can evoke a resonance in the observer's motor system. The observer's 'motor knowledge' is harnessed to decipher the observed action. However, mirror neurons extend beyond being merely neural foundations for our sporting inclinations. They also enable observers to enhance their athletic abilities through the process of observation. It turns out that watching a remarkable golfer improves our own golfing skills, and witnessing a skilled sprinter actually bolsters our running speed. This ability to learn through observation is a vital skill. From language acquisition in infancy to understanding facial expressions, mimicking plays a fundamental role in our consciousness. The most accomplished athletes possess a premotor cortex capable of envisioning victorious movements, supplemented by the physical attributes required to translate those movements into reality. In this way, action and imagination coalesce to drive performance.")
p("F However, how many among us watch sport in order to enhance our athletic abilities? Most spectators tune in for the emotional roller coaster, the human drama it encompasses. This emotional connection finds its roots in mirror neurons, which allow spectators to perceive athletes' victorious movements and to share in the emotion of their triumphs. This is because they are directly associated with the amygdala, a key area of the brain responsible for emotional processing. During the Olympics, mirror neurons in individuals around the world resonate in unison, and watching sport unifies people. The majority of us may never run a mile in under four minutes or hit a home run. Yet, as we gather around our televisions, we all briefly get a taste of the feeling of executing something flawlessly. In those fleeting moments, sport becomes more than entertainment-it becomes a shared human experience.")

# QUESTIONS - Passage 3
p("[QUESTIONS]")

# Q27-32: Which paragraph (MATCHING)
p("[MATCHING]")
p("Questions 27-32")
p("Reading Passage 3 has six paragraphs, A-F.")
p("Which paragraph contains the following information?")
p("Choose the correct letter, A-F, in boxes 27-32.")
p("NB You may use any letter more than once.")
p("A")
p("B")
p("C")
p("D")
p("E")
p("F")

p("27. reasons why watching sport can bring emotional satisfaction")
p("*F")

p("28. an explanation of why beginners struggle with sports tasks")
p("*B")

p("29. a factor that can work alongside mirroring in sports competitions to maintain peak performance")
p("*E")

p("30. a comparison between the mirror neurons of human beings and primates")
p("*C")

p("31. mention of integrating visual and motor processes to understand athletic movements")
p("*E")

p("32. reference to the first discovery of brain activity associated with mirror neurons")
p("*D")

# Q33-35: Multiple choice (CHOICE)
p("[CHOICE]")
p("Questions 33-35")
p("Choose the correct answer.")

p("33. What point does the writer make about professional athletes in Paragraph B?")
p("A. They meticulously plan each phase of their movements.")
p("B. They engage in regular practice to retain and refine their skills.")
p("C. They actively consider and analyse the actions of their peers.")
p("D. They perform their actions without conscious contemplation.")
p("*D")

p("34. In this passage, the writer mentions the example of a remarkable golfer to illustrate")
p("A. the critical role of imitation in the formation of consciousness.")
p("B. the improvement of sporting ability through watching sport.")
p("C. the stimulation of a resonance in the observer's motor system.")
p("D. the motivation of the observer's interest and love for golf.")
p("*B")

p("35. The principal motivation for watching sport mentioned by the writer is")
p("A. to improve our physical ability by analyzing the athletes' technical strategies")
p("B. to unite viewers from different nations.")
p("C. to experience intense positive emotions by immersing oneself in sport.")
p("D. to become a sports professional.")
p("*C")

# Q36-40: YES/NO/NOT GIVEN (CHOICE)
p("[CHOICE]")
p("Questions 36-40")
p("Choose YES if the statement agrees with the claims of the writer, choose NO if the statement contradicts the claims of the writer, or NOT GIVEN if it is impossible to say what the writer thinks about this. Yes / No / Not Given")

p("36. The emergence of Homer's epic poetry occurred in the same historical period as the beginning of the Olympic Games.")
p("*YES")

p("37. The Italian scientists concluded that the motor actions encoded within the premotor cortex have no discernible link to unrelated collections of neurons.")
p("*NOT GIVEN")

p("38. Spectators can achieve flawless golf swings merely by watching the Masters Tournament without physical practice.")
p("*NO")

p("39. Some neuroscientists believe that certain types of still images might activate mirror neurons under specific experimental conditions.")
p("*NO")

p("40. The mu system is suppressed while we are dynamically involved in a particular activity.")
p("*YES")

import sys
out = sys.argv[1] if len(sys.argv) > 1 else "Test 3.docx"
doc.save(out)
print("[OK] da tao xong:", out)
