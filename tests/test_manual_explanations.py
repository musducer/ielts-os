import unittest
from unittest.mock import patch

import docx

from api.index import (
    _AI_KEY_CURSOR,
    _api_key_pool,
    _rotated_api_keys,
    extract_docx_manual_explanations,
    parse_docx_to_quiz,
    parse_manual_explanation_text,
)


class ManualExplanationTests(unittest.TestCase):
    def test_manual_explanation_parser_extracts_quotes_and_timestamps(self):
        parsed = parse_manual_explanation_text(
            'Because the speaker says "annual membership fee" at [02:15 - 02:20].'
        )

        self.assertEqual(parsed["rawText"], 'Because the speaker says "annual membership fee" at [02:15 - 02:20].')
        self.assertEqual(parsed["parsedQuotes"], ["annual membership fee"])
        self.assertEqual(parsed["parsedTimestamps"], [
            {"startTime": 135, "label": "02:15 - 02:20", "endTime": 140},
        ])

    def test_docx_parser_attaches_explanation_without_rendering_tag(self):
        document = docx.Document()
        document.add_paragraph("[TITLE] Manual Explanation Demo")
        document.add_paragraph("[TYPE] Listening")
        document.add_paragraph("[PASSAGE]")
        document.add_paragraph("The transcript contains the annual membership fee.")
        document.add_paragraph("[QUESTIONS]")
        document.add_paragraph("[BLANK]")
        document.add_paragraph("Question 1")
        document.add_paragraph("*annual membership fee")
        document.add_paragraph('[EXPLANATION] 1: The speaker says "annual membership fee" at [02:15].')

        quiz = parse_docx_to_quiz(document)
        question = quiz["questions"][0]

        self.assertNotIn("[EXPLANATION]", question["text"])
        self.assertEqual(question["manualExplanation"]["rawText"], 'The speaker says "annual membership fee" at [02:15].')
        self.assertEqual(question["manualExplanation"]["parsedQuotes"], ["annual membership fee"])
        self.assertEqual(question["manualExplanation"]["parsedTimestamps"], [
            {"startTime": 135, "label": "02:15"},
        ])

    def test_matching_answer_is_stored_as_text_not_letter_or_index(self):
        document = docx.Document()
        document.add_paragraph("[TITLE] Matching Demo")
        document.add_paragraph("[TYPE] Listening")
        document.add_paragraph("[PASSAGE]")
        document.add_paragraph("Transcript.")
        document.add_paragraph("[QUESTIONS]")
        document.add_paragraph("[MATCHING]")
        document.add_paragraph("A. Processing data she had gathered")
        document.add_paragraph("B. Working collaboratively")
        document.add_paragraph("Question 26 Observing lessons")
        document.add_paragraph("*A")

        quiz = parse_docx_to_quiz(document)
        question = quiz["questions"][0]

        self.assertEqual(question["type"], "MATCHING")
        self.assertEqual(question["options"], [
            "Processing data she had gathered",
            "Working collaboratively",
        ])
        self.assertEqual(question["correctAnswer"], "Processing data she had gathered")

    def test_explanation_only_docx_supports_multiline_entries(self):
        document = docx.Document()
        document.add_paragraph('[EXPLANATION] 26: The speaker says "processing data" at [12:04].')
        document.add_paragraph("[EXPLANATION] 27:")
        document.add_paragraph('The clue is "suitable equipment" at [12:10 - 12:13].')

        explanations = extract_docx_manual_explanations(document)

        self.assertEqual(explanations["26"]["parsedQuotes"], ["processing data"])
        self.assertEqual(explanations["27"]["rawText"], 'The clue is "suitable equipment" at [12:10 - 12:13].')
        self.assertEqual(explanations["27"]["parsedTimestamps"], [
            {"startTime": 730, "label": "12:10 - 12:13", "endTime": 733},
        ])

    def test_api_key_pool_supports_numbered_and_comma_separated_keys(self):
        with patch.dict("os.environ", {
            "GROQ_API_KEY": "g1",
            "GROQ_API_KEY_2": "g2",
            "GROQ_API_KEYS": "g3,g4",
        }, clear=False):
            keys = _api_key_pool("GROQ_API_KEY")

        self.assertEqual(keys[:4], ["g1", "g2", "g3", "g4"])

    def test_rotated_api_keys_distributes_starting_key(self):
        _AI_KEY_CURSOR.clear()

        self.assertEqual(_rotated_api_keys("test_provider", ["a", "b", "c"]), ["a", "b", "c"])
        self.assertEqual(_rotated_api_keys("test_provider", ["a", "b", "c"]), ["b", "c", "a"])
        self.assertEqual(_rotated_api_keys("test_provider", ["a", "b", "c"]), ["c", "a", "b"])


if __name__ == "__main__":
    unittest.main()
