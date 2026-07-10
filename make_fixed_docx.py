# -*- coding: utf-8 -*-
"""
Tạo bản .docx CHUẨN HOÁ + ĐỊNH DẠNG ĐẸP cho đề READING TEST 2 - NGỌC ÁNH.
Quy tắc định dạng (theo yêu cầu):
  - Notes-completion (Q9-13): item -> bullet list.
  - Tiêu đề bài đọc (Heading 2) -> IN ĐẬM + CĂN GIỮA.
  - Tiêu đề flow-chart / notes / summary -> IN ĐẬM + CĂN GIỮA.
  - Sa-pô (đoạn dẫn dưới tiêu đề, nếu có) -> IN NGHIÊNG + CĂN GIỮA.
  - Heading phụ trong notes (Japan / Other countries) -> in đậm.
Giữ nguyên toàn bộ nội dung/đáp án.
"""
import os, docx
from docx.enum.text import WD_ALIGN_PARAGRAPH

SRC = r"C:\Users\hp\Downloads\UPGRADED_READING_TEST_2_NGOC_ANH.docx"
OUT = r"C:\Users\hp\Downloads\READING_TEST_2_NGOC_ANH_FIXED.docx"

d = docx.Document(SRC)

def set_bullet(p):
    for name in ("List Bullet", "ListBullet", "List Paragraph"):
        try:
            p.style = d.styles[name]; return True
        except KeyError:
            continue
    if p.runs and not p.runs[0].text.lstrip().startswith("•"):
        p.runs[0].text = "•  " + p.runs[0].text
    return False

def bold(p):
    if not p.runs and p.text:
        p.add_run(p.text)
    for r in p.runs: r.bold = True

def italic(p):
    for r in p.runs: r.italic = True

def center(p):
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

NOTE_ITEMS   = ["[9]", "[10]", "[11]", "[12]", "[13]"]
NOTE_SUBHEAD = {"Japan", "Other countries"}
# Tiêu đề block (flow-chart / notes / summary) -> bold + center
BLOCK_TITLES = {
    "Building the Wood Innovation and Design Centre",
    "Other Significant Wooden Buildings",
    "The Inter-Agency Space Debris Coordination Committee",
}
# Sa-pô (deck) -> italic + center
STANDFIRST_PREFIXES = (
    "Satellites, rocket shards and collision debris",
)

cnt = {"bullet": 0, "title": 0, "block": 0, "deck": 0, "sub": 0}
for p in d.paragraphs:
    t = p.text.strip()
    if not t:
        continue
    style = (p.style.name or "").lower() if p.style else ""

    # 1) Tiêu đề bài đọc (Heading 2) -> bold + center
    if "heading 2" in style:
        bold(p); center(p); cnt["title"] += 1; continue
    # 2) Tiêu đề block completion -> bold + center
    if t in BLOCK_TITLES:
        bold(p); center(p); cnt["block"] += 1; continue
    # 3) Sa-pô -> italic + center
    if any(t.startswith(pre) for pre in STANDFIRST_PREFIXES):
        italic(p); center(p); cnt["deck"] += 1; continue
    # 4) Notes: item -> bullet
    if any(m in t for m in NOTE_ITEMS):
        set_bullet(p); cnt["bullet"] += 1; continue
    # 5) Notes sub-heading -> bold
    if t in NOTE_SUBHEAD:
        bold(p); cnt["sub"] += 1; continue

d.save(OUT)
print(f"[OK] {OUT}")
print("     ", cnt, "| size:", os.path.getsize(OUT), "bytes")
