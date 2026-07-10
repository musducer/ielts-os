# -*- coding: utf-8 -*-
"""
Sinh 1 file .docx CHUẨN FORMAT BACKEND (parse_docx_to_quiz) cho:
  VOL 9 - TEST 4 LISTENING  (đề + đáp án gộp làm một)
Upload qua nút "Upload .docx" trong Exam Builder -> ra quiz đủ 40 câu, 4 section, đủ đáp án.
"""
from docx import Document

doc = Document()

def p(text="", bold=False):
    para = doc.add_paragraph()
    r = para.add_run(text)
    if bold:
        r.bold = True
    return para

def tbl(rows):
    t = doc.add_table(rows=len(rows), cols=len(rows[0]))
    t.style = "Table Grid"
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            t.cell(ri, ci).text = val
    return t

# ================= META =================
p("[TITLE] VOL 9 - Test 4 (Listening)")
p("[TIME] 40")
p("[TYPE] Listening")
p("[AUDIO] ")   # <-- Dán link audio TEST 4 vào đây (hoặc thêm trong Builder sau khi upload)

# ============================================================
# SECTION 1 (Q1-10)  — FORM COMPLETION
# ============================================================
p("[PASSAGE]")
p("[QUESTIONS]")
p("[SHORT_ANSWER]")
p("Questions 1–10")
p("Complete the form below.")
p("Write ONE WORD AND/OR A NUMBER for each answer.")

p("[CONTEXT]")
p("Weekend Course Booking", bold=True)
p("Booking for: a guitar workshop")
p("Start date: [1]")
p("Course code: [2]")
p("Previous course attended: Editing [3]")
p("Name: Nick [4]")
p("Address: 27 Hawthorn [5] Bristol, BS4 2QG")
p("Contact number: 07707 438 297")
p("Age: 26")
p("Occupation: [6]")
p("Dietary requests: no [7]")
p("One-to-one lesson selected: [8] guitar")
p("First heard about us from: a [9]")
p("Additional request: would like [10] if possible")
p("[/CONTEXT]")

for n, ans in [(1,"July 18th"),(2,"Q1632"),(3,"photographs"),(4,"Poskitt"),
               (5,"way"),(6,"engineer"),(7,"fish"),(8,"blues"),
               (9,"magazine"),(10,"parking")]:
    p(f"{n}.")
    p(f"*{ans}")

# ============================================================
# SECTION 2 (Q11-20)
# ============================================================
p("[PASSAGE]")
p("[QUESTIONS]")

# Q11-15: Matching A-G -> KÉO-THẢ (mỗi đáp án dùng 1 lần) => dùng [DRAG]
p("[DRAG]")
p("Questions 11–15")
p("What is the main attraction in each of the following areas of the festival?")
p("Choose FIVE correct answers, A–G, and move each into the gap next to questions 11–15.")
p("Attractions")
for opt in ["A. talks about gardening","B. live music","C. an art exhibition",
            "D. a children's playground","E. evening entertainment",
            "F. a competition for adults","G. vegetable market"]:
    p(opt)
p("Areas of the festival")
for n, area, ans in [(11,"Bridge Street","C"),(12,"King's Park","D"),
                     (13,"Oakvale School","F"),(14,"Green Road","B"),
                     (15,"Sun Lane","A")]:
    p(f"{n}. {area}")
    p(f"*{ans}")

# Q16-20: Multiple choice
p("[CHOICE]")
p("Questions 16–20")
p("Choose the correct answer.")
p("Denford Festival")

mcq_2 = [
    (16,"What does Sam say about the festival tickets?",
     ["They allow you to get into all the festival events.",
      "They are more expensive to buy on the day.",
      "They have nearly all been sold."], "B"),
    (17,"On the last night of the festival, visitors will be able to",
     ["take part in a dance.","go to a theatre performance.","watch some fireworks."], "A"),
    (18,"What will be provided for international visitors?",
     ["discounts on accommodation.","tours of the local area.",
      "a guide written in different languages."], "B"),
    (19,"What information is given about car park A?",
     ["It is bigger than car park B.","It is closer to town than car park B.",
      "It will close earlier than car park B."], "A"),
    (20,"What change have the organisers made to the programme this year?",
     ["replacing the programme with an app","reducing the number of programmes printed",
      "introducing a charge for the programmes"], "B"),
]
for n, stem, opts, ans in mcq_2:
    p(f"{n}. {stem}")
    for i, o in enumerate(opts):
        p(f"{chr(65+i)}. {o}")
    p(f"*{ans}")

# ============================================================
# SECTION 3 (Q21-30)
# ============================================================
p("[PASSAGE]")
p("[QUESTIONS]")

# Q21-23: Note completion
p("[SHORT_ANSWER]")
p("Questions 21–23")
p("Complete the notes below.")
p("Write NO MORE THAN TWO WORDS for each answer.")
p("[CONTEXT]")
p("Action Research Course for Practising Teachers", bold=True)
p("Course leader: Alan Cumner")
p("    (get his recent book called '[21]')")
p("Course objectives: to develop skills in:")
p("•  using different research techniques")
p("•  researching as part of a [22]")
p("•  [23] (in a collaborative culture)")
p("[/CONTEXT]")
for n, ans in [(21,"professional learning"),(22,"team"),(23,"presenting results")]:
    p(f"{n}.")
    p(f"*{ans}")

# Q24-30: Table completion (2 tables)
p("[SHORT_ANSWER]")
p("Questions 24–30")
p("Complete the tables below.")
p("Write NO MORE THAN TWO WORDS for each answer.")
p("Course Components")
p("[CONTEXT]")
tbl([
    ["Observational techniques","Additional information"],
    ["Using observation checklists","To record aspects of lesson, e.g. [24] of pupils"],
    ["Writing a [25]","Used to improve skill of reflection"],
    ["[26]","Try out on fellow students (in classroom [27])"],
])
tbl([
    ["Non-observational techniques","Additional information"],
    ["Analysing [28]","Use statistics based on own students"],
    ["Using questionnaires","Use the [29] to find respondents"],
    ["Using [30]","Own choice of respondents"],
])
p("[/CONTEXT]")
for n, ans in [(24,"behaviour"),(25,"diary"),(26,"video recording"),
               (27,"simulation"),(28,"test results"),(29,"internet"),(30,"interviews")]:
    p(f"{n}.")
    p(f"*{ans}")

# ============================================================
# SECTION 4 (Q31-40)
# ============================================================
p("[PASSAGE]")
p("[QUESTIONS]")

# Q31-33: Multiple choice
p("[CHOICE]")
p("Questions 31–33")
p("Choose the correct answer.")
mcq_4 = [
    (31,"What impact does Marc Prensky believe that digital technology has had on young people?",
     ["It has altered their thinking patterns.","It has harmed their physical development.",
      "It has limited their brain capacity."], "A"),
    (32,"Digital immigrants tend to access computers",
     ["using their native language.","less efficiently than young people.",
      "for less important information."], "B"),
    (33,"What example is given of having a 'digital accent'?",
     ["having less effective typing skills","doing things the old-fashioned way",
      "being unable to understand instructions"], "B"),
]
for n, stem, opts, ans in mcq_4:
    p(f"{n}. {stem}")
    for i, o in enumerate(opts):
        p(f"{chr(65+i)}. {o}")
    p(f"*{ans}")

# Q34-40: Matching theorists A/B/C -> KÉO-THẢ tái dùng (7 câu / 3 đáp án) => [DRAG]
p("[DRAG]")
p("Questions 34–40")
p("Which theorist makes each of the following points?")
p("Choose the correct answer for each point and move it into the gap. You may use any theorist more than once.")
p("Theorists")
for opt in ["A. Allen","B. James","C. Vander"]:
    p(opt)
p("Points made")
for n, point, ans in [
    (34,"Current teaching methods don't work.","B"),
    (35,"Many students don't understand computers.","C"),
    (36,"Computer technology doesn't interest all students.","A"),
    (37,"Students can still learn the traditional way.","B"),
    (38,"Students still need to learn research skills.","A"),
    (39,"We should use computer games to teach.","B"),
    (40,"Computers can't replace educators.","C"),
]:
    p(f"{n}. {point}")
    p(f"*{ans}")

import sys
out = sys.argv[1] if len(sys.argv) > 1 else "VOL9 TEST4 - Listening (IELTS OS upload).docx"
doc.save(out)
print("[OK] Saved:", out)
