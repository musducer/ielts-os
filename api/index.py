import re
import io
import os
import time
import traceback
import base64
import urllib.parse
from fastapi import FastAPI, UploadFile, File, Body, Request
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import JSONResponse, RedirectResponse
import docx
from docx.table import Table
from docx.text.paragraph import Paragraph
from typing import List, Dict, Any

app = FastAPI(docs_url="/api/docs", openapi_url="/api/openapi.json")
app.add_middleware(CORSMiddleware, allow_origins=["*"], allow_credentials=True, allow_methods=["*"], allow_headers=["*"])


@app.exception_handler(Exception)
async def unexpected_api_error(request: Request, exc: Exception):
    """Keep API failures JSON-shaped so the client can show a useful recovery message."""
    print(f"Unhandled API error for {request.url.path}: {exc}")
    traceback.print_exc()
    return JSONResponse(status_code=500, content={
        "success": False,
        "error": "Máy chủ AI gặp lỗi khi xử lý yêu cầu. Vui lòng thử lại sau ít phút.",
    })

FIREBASE_STORAGE_BUCKET = os.environ.get("FIREBASE_STORAGE_BUCKET", "ielts-os.firebasestorage.app")

VALID_BLOCK_TYPES = ["BLANK", "CHOICE", "CHOICE_MULTIPLE", "MATCHING", "DRAG_DROP", "DRAG", "SHORT_ANSWER"]

RE_Q_START = re.compile(r'^(?:Question|Câu)\s+\d+', re.IGNORECASE)
RE_Q_NUMBER = re.compile(r'^\d+[\.\)\:]')
RE_TFNG = re.compile(r'\btrue\b[\s\S]*\bfalse\b[\s\S]*\bnot\s+given\b', re.IGNORECASE)
RE_YNNG = re.compile(r'\byes\b[\s\S]*\bno\b[\s\S]*\bnot\s+given\b', re.IGNORECASE)

def is_option(text: str) -> bool:
    text = text.strip()
    # Dạng 1 chuẩn: A. text, A) text
    if re.match(r'^[A-Za-z][\.\)]\s+', text): return True
    # Dạng 2: A, B, C đứng độc lập
    if re.match(r'^[A-Za-z]$', text): return True
    # Dạng roman numeral heading: i. text, ii. text, iii. text
    if re.match(r'^[ivxlcdmIVXLCDM]+[\.]\)\s+\S', text): return True
    if re.match(r'^[ivxlcdmIVXLCDM]+\.\s+\S', text): return True  # roman numeral heading
    # Dạng 3: A Text (List of people, features...)
    m = re.match(r'^([A-K])\s+(.+)$', text)
    if m:
        letter, rest = m.groups()
        if len(text) < 80: # Giới hạn độ dài để tránh nhận nhầm câu văn bản dài
            if letter == 'A' and rest[0].islower(): return False # "A boy walked..."
            if letter == 'I' and (rest.lower().startswith('am ') or rest.lower().startswith('have ') or rest.lower().startswith('was ')): return False
            return True
    return False

def paragraph_to_html(para: Paragraph) -> str:
    if not para.text: return ""
    html_parts = []
    for run in para.runs:
        text = run.text or ""
        if not text: continue
        text = text.replace('\t', '&nbsp;&nbsp;&nbsp;&nbsp;').replace('  ', '&nbsp;&nbsp;').replace('\n', '<br/>')
        if run.bold: text = f"<strong>{text}</strong>"
        if run.italic: text = f"<em>{text}</em>"
        if run.underline: text = f"<u>{text}</u>"
        html_parts.append(text)
    p_html = "".join(html_parts)
    if not p_html.replace('&nbsp;', '').strip() and '<br/>' not in p_html:
        return ""
    
    align = para.alignment
    style_name = para.style.name.lower() if para.style else ""
    wrapper_style = ""
    if align == 1: wrapper_style += "text-align: center; "
    elif align == 2: wrapper_style += "text-align: right; "
    elif align == 3: wrapper_style += "text-align: justify; "
    
    # BẢO TỒN ĐỊNH DẠNG BULLET POINTS CHO NOTES COMPLETION
    is_list = False
    if 'list' in style_name or 'bullet' in style_name:
        is_list = True
    elif para._p.pPr is not None and para._p.pPr.numPr is not None:
        is_list = True

    if 'heading 1' in style_name: return f'<h1 style="{wrapper_style}margin:15px 0;font-size:24px;">{p_html}</h1>'
    if 'heading 2' in style_name: return f'<h2 style="{wrapper_style}margin:12px 0;font-size:20px;">{p_html}</h2>'
    if 'heading 3' in style_name: return f'<h3 style="{wrapper_style}margin:10px 0;font-size:18px;">{p_html}</h3>'
    
    if is_list:
        return f'<div style="display: list-item; margin-left: 24px; list-style-type: disc; {wrapper_style}">{p_html}</div>'

    return f'<div style="{wrapper_style}">{p_html}</div>' if wrapper_style else p_html

def table_to_html(table: Table) -> str:
    rows = table.rows
    if not rows: return ""
    n_rows, n_cols = len(rows), len(rows[0].cells)
    html = '<table style="border-collapse:collapse; margin-bottom:15px; table-layout:auto;"><tbody>'
    for r_idx in range(n_rows):
        html += "<tr>"
        for c_idx in range(n_cols):
            raw_text = " ".join([p.text.strip() for p in table.cell(r_idx, c_idx).paragraphs if p.text.strip()])
            cell_text = "".join([f"<div>{paragraph_to_html(p)}</div>" for p in table.cell(r_idx, c_idx).paragraphs if p.text.strip()])
            is_short = len(raw_text) <= 20
            nowrap = "white-space:nowrap;" if is_short else ""
            is_header = r_idx == 0
            bg = "background:#f5f5f5; font-weight:bold;" if is_header else ""
            html += f'<td style="border:1px solid #ccc; padding:6px 10px; {nowrap}{bg} vertical-align:top;">{cell_text}</td>'
        html += "</tr>"
    html += "</tbody></table>"
    return html

def process_block(block_type: str, lines: List[Any], target_questions: List[Dict]):
    html_lines, text_lines = [], []
    for item in lines:
        if isinstance(item, Paragraph):
            h = paragraph_to_html(item)
            if not h: continue
            html_lines.append(h)
            text_lines.append(re.sub(r'<[^>]+>', '', h).strip())
        elif isinstance(item, Table):
            html_lines.append(table_to_html(item))
            text_lines.append("[TABLE]")
    
    questions_in_block = []
    current_q = None
    accumulated_instruction = ""
    group_context = ""
    pre_question_options = []
    in_context = False
    
    for h_line, t_line in zip(html_lines, text_lines):
        t_upper = t_line.upper()
        if "[CONTEXT]" in t_upper:
            in_context = True
            continue
        if "[/CONTEXT]" in t_upper:
            in_context = False
            continue
        
        if in_context:
            group_context += f"<div>{h_line}</div>"
            continue

        is_q_start = bool(RE_Q_START.match(t_line) or RE_Q_NUMBER.match(t_line))
        
        if not is_q_start and not current_q:
            # Matching-section: nhãn "A\nB\nC\nD\nE\nF" bị gộp 1 đoạn -> tách thành từng option
            parts = [re.sub(r'<[^>]+>', '', p).strip() for p in re.split(r'<br\s*/?>', h_line)]
            parts = [p for p in parts if p]
            if len(parts) >= 2 and all(re.fullmatch(r'[A-Za-z]', p) for p in parts):
                pre_question_options.extend([p.upper() for p in parts])
                continue
            if is_option(t_line):
                pre_question_options.append(t_line)
                continue
            accumulated_instruction += f"<div>{h_line}</div>"
            continue

        if is_q_start:
            if current_q: questions_in_block.append(current_q)
            content_only = re.sub(r'^(?:Question|Câu)?\s*\d+[\.\:\)]?\s*', '', t_line, flags=re.IGNORECASE).strip()
            current_q = {
                "type": block_type, "text": content_only if content_only else h_line,
                "options": pre_question_options.copy(), "correctAnswers": [],
                "instruction": accumulated_instruction,
                "groupContext": group_context
            }
            continue
        
        if current_q:
            if t_line.startswith("*"):
                current_q["correctAnswers"].append(re.sub(r'^\*+\s*', '', t_line).strip())
            elif is_option(t_line):
                current_q["options"].append(t_line)
            else:
                current_q["text"] += f" <div>{h_line}</div>"

    if current_q: questions_in_block.append(current_q)
    
    # Chia sẻ options cho các câu trong nhóm
    shared_options = []
    for q in questions_in_block:
        if q["options"]: shared_options = q["options"]
        elif shared_options and q["type"] in ["CHOICE", "CHOICE_MULTIPLE", "MATCHING", "DRAG"]:
            q["options"] = shared_options.copy()
    
    for q in questions_in_block:
        lower_txt = (q['instruction'] + q['text']).lower()
        sub_t = "SENTENCE"
        if "flow-chart" in lower_txt or "flowchart" in lower_txt: sub_t = "FLOWCHART"
        elif "summary" in lower_txt: sub_t = "SUMMARY"
        elif "note" in lower_txt: sub_t = "NOTES"
        
        is_tfng = RE_TFNG.search(lower_txt) or "true/false/not given" in lower_txt
        is_ynng = RE_YNNG.search(lower_txt) or "yes/no/not given" in lower_txt
        
        if not q["options"]:
            if is_tfng: q["options"] = ["TRUE", "FALSE", "NOT GIVEN"]
            elif is_ynng: q["options"] = ["YES", "NO", "NOT GIVEN"]
            
        output_type = "DRAG_DROP_HEADING" if "heading" in lower_txt and q["type"] == "MATCHING" else q["type"]
        if is_tfng or is_ynng: output_type = "CHOICE"
        # SYNC FRONTEND: Short answer = điền từ -> render như Inline Blank để FE có ô nhập
        if output_type == "SHORT_ANSWER": output_type = "BLANK"
        # WORD-BANK (tóm tắt kéo-thả từ): block [DRAG] -> DRAG_DROP, options là CÁC TỪ (bỏ tiền tố chữ cái A-K)
        if q["type"] == "DRAG":
            output_type = "DRAG_DROP"
            q["options"] = [re.sub(r'^\s*[A-Za-z][\.\)]\s*', '', str(o)).strip() for o in (q["options"] or [])]

        final_correct_answer = ""
        # THUẬT TOÁN ÁNH XẠ ĐÁP ÁN (SMART MAPPER)
        if output_type in ["CHOICE", "MATCHING"]:
            if q["correctAnswers"] and q["options"]:
                ans_str = q["correctAnswers"][0].strip().upper()
                for idx, opt in enumerate(q["options"]):
                    opt_str = opt.strip()
                    m_letter = re.match(r'^([A-Za-z])(?:[\.\)]\s+|\s+|$)', opt_str)
                    opt_letter = m_letter.group(1).upper() if m_letter else chr(65 + idx)
                    
                    if ans_str == opt_letter or ans_str == opt_str.upper():
                        final_correct_answer = idx
                        break
                if final_correct_answer == "": final_correct_answer = 0
            else:
                final_correct_answer = 0
        elif output_type == "CHOICE_MULTIPLE":
            ans_arr = []
            if q["correctAnswers"]:
                for ans in q["correctAnswers"]:
                    ans_str = ans.strip().upper()
                    for idx, opt in enumerate(q["options"] if q["options"] else []):
                        opt_str = opt.strip()
                        m_letter = re.match(r'^([A-Za-z])(?:[\.\)]\s+|\s+|$)', opt_str)
                        opt_letter = m_letter.group(1).upper() if m_letter else chr(65 + idx)
                        if ans_str == opt_letter or ans_str == opt_str.upper():
                            ans_arr.append(idx)
                            break
            final_correct_answer = ans_arr
        elif output_type == "DRAG_DROP_HEADING":
            # correctAnswer = roman numeral label (e.g. "iv") matched against option prefix
            if q["correctAnswers"]:
                raw = str(q["correctAnswers"][0]).strip().lower()
                final_correct_answer = raw  # default: giữ nguyên
                for opt in (q["options"] or []):
                    m = re.match(r'^([ivxlcdmIVXLCDM]+)[\.)\s]', str(opt).strip(), re.IGNORECASE)
                    if m and m.group(1).lower() == raw:
                        final_correct_answer = m.group(1).lower()
                        break
            else:
                final_correct_answer = ""
        elif output_type == "DRAG_DROP":
            # Đáp án có thể là chữ cái (A-K) -> ánh xạ sang TỪ tương ứng (vì chấm điểm DRAG_DROP so khớp text);
            # nếu đã là chữ thuần thì giữ nguyên.
            if q["correctAnswers"]:
                raw = str(q["correctAnswers"][0]).strip()
                if re.fullmatch(r'[A-Za-z]', raw) and q["options"]:
                    idx = ord(raw.upper()) - 65
                    final_correct_answer = q["options"][idx] if 0 <= idx < len(q["options"]) else raw
                else:
                    final_correct_answer = raw
            else:
                final_correct_answer = ""
        else:
            final_correct_answer = q["correctAnswers"][0] if q["correctAnswers"] else ""

        target_questions.append({
            "id": f"q_{int(time.time() * 1000)}_{len(target_questions)}",
            "type": output_type,
            "subType": sub_t, 
            "instruction": q["instruction"], 
            "groupContext": q["groupContext"],
            "text": q["text"], 
            "options": q["options"] if q["options"] else [],
            "correctAnswer": final_correct_answer
        })

def parse_docx_to_quiz(doc):
    title, time_limit, quiz_type, audio_url = "Untitled Mock Test", 60, "Reading", ""
    sections = []
    
    current_passage = ""
    current_questions = []
    
    state = "META"
    block_type = None
    current_block_lines = []

    def flush_block():
        nonlocal current_block_lines, block_type
        if block_type and current_block_lines:
            process_block(block_type, current_block_lines, current_questions)
        current_block_lines = []

    for element in doc.element.body:
        if element.tag.endswith('p'):
            para = Paragraph(element, doc)
            text = para.text.strip()
            if not text: continue
            text_upper = text.upper()
            
            if "[TITLE]" in text_upper: title = re.sub(r'\[TITLE\]', '', text, flags=re.IGNORECASE).strip(); continue
            if "[TIME]" in text_upper: 
                try: time_limit = int(re.sub(r'\D', '', text))
                except: pass
                continue
            if "[TYPE]" in text_upper: quiz_type = re.sub(r'\[TYPE\]', '', text, flags=re.IGNORECASE).strip(); continue
            if "[AUDIO]" in text_upper: audio_url = re.sub(r'\[AUDIO\]', '', text, flags=re.IGNORECASE).strip(); continue
            
            if "[PASSAGE]" in text_upper:
                flush_block()
                if current_passage or current_questions:
                    sections.append({"passage": current_passage, "questions": current_questions})
                current_passage, current_questions = "", []
                state = "PASSAGE"
                block_type = None
                continue
            
            if "[QUESTIONS]" in text_upper:
                flush_block()
                state = "QUESTIONS"
                block_type = None
                continue

            if state == "PASSAGE":
                html = paragraph_to_html(para)
                if html: current_passage += f"<div style='margin-bottom:12px;'>{html}</div>"
            elif state == "QUESTIONS":
                match = re.search(r'\[(/?\w+)\]', text_upper)
                if match and match.group(1) in VALID_BLOCK_TYPES:
                    flush_block()
                    block_type = match.group(1)
                    continue
                
                if block_type:
                    current_block_lines.append(para)
                
        elif element.tag.endswith('tbl'):
            if state == "PASSAGE":
                current_passage += table_to_html(Table(element, doc))
            elif state == "QUESTIONS" and block_type:
                current_block_lines.append(Table(element, doc))

    flush_block()
    if current_passage or current_questions:
        sections.append({"passage": current_passage, "questions": current_questions})

    # Gắn passageIndex để FE giữ 1 nguồn chân lý (mảng phẳng) và ánh xạ đúng section.
    for si, sec in enumerate(sections):
        for q in sec["questions"]:
            q["passageIndex"] = si

    all_questions = []
    for sec in sections: all_questions.extend(sec["questions"])

    return {
        "id": f"quiz_{int(time.time() * 1000)}",
        "title": title,
        "type": quiz_type,
        "audioUrl": audio_url,
        "timeLimit": time_limit,
        "passage": sections[0]["passage"] if sections else "",
        "sections": sections,
        "questions": all_questions,
        "active": False,
        "maxAttempts": 1,
        "audience": "ALL",
        "targetStudentIds": [],
        "folder": "Root",
        "isSEBRequired": False,
        "passcode": "",
        "scheduledStart": "",
        "scheduledEnd": ""
    }

@app.post("/api/upload_docx")
async def upload_docx(file: UploadFile = File(...)):
    try:
        doc = docx.Document(io.BytesIO(await file.read()))
        quiz = parse_docx_to_quiz(doc)
        if not quiz.get("questions") or len(quiz["questions"]) == 0:
            return {"success": False, "error": "Lỗi: Không trích xuất được câu hỏi nào. Hãy chắc chắn bạn đã gắn đủ thẻ [PASSAGE] và [QUESTIONS]."}
        return {"success": True, "quiz": quiz}
    except Exception as e:
        return {"success": False, "error": f"Lỗi xử lý file: {str(e)}\n{traceback.format_exc()}"}

def _decode_audio_key(audio_key: str):
    raw = audio_key.replace("-", "+").replace("_", "/")
    raw += "=" * (-len(raw) % 4)
    decoded = base64.b64decode(raw.encode("ascii")).decode("utf-8")
    storage_path, token = decoded.split("|", 1)
    if not storage_path.startswith("exam-audio/") or not token:
        raise ValueError("invalid audio key")
    return storage_path, token

@app.get("/api/audio/{audio_key}/{filename}")
async def hosted_audio(audio_key: str, filename: str):
    try:
        storage_path, token = _decode_audio_key(audio_key)
        target = (
            "https://firebasestorage.googleapis.com/v0/b/"
            + urllib.parse.quote(FIREBASE_STORAGE_BUCKET, safe="")
            + "/o/"
            + urllib.parse.quote(storage_path, safe="")
            + "?alt=media&token="
            + urllib.parse.quote(token, safe="")
        )
        return RedirectResponse(target, status_code=302)
    except Exception:
        return {"success": False, "error": "Audio link không hợp lệ hoặc đã bị đổi."}

def _oai_chat_once(base_url, api_key, model, sys_prompt, user_prompt, max_tokens, temperature, json_mode, reasoning_effort):
    """1 lần gọi API OpenAI-compatible (Groq/Cerebras). Trả (text, err). err: 'RATE_LIMIT'/'AUTH'/chi tiết."""
    import json as _json
    import urllib.request as _urlreq
    import urllib.error as _urlerr
    body = {
        "model": model,
        "messages": [{"role": "system", "content": sys_prompt}, {"role": "user", "content": user_prompt}],
        "max_tokens": max_tokens, "temperature": temperature,
    }
    if json_mode:
        body["response_format"] = {"type": "json_object"}
    if reasoning_effort and "gpt-oss" in model:
        body["reasoning_effort"] = reasoning_effort
    try:
        req = _urlreq.Request(
            base_url.rstrip("/") + "/chat/completions",
            data=_json.dumps(body).encode("utf-8"), method="POST",
            headers={"Authorization": f"Bearer {api_key}", "Content-Type": "application/json",
                     "Accept": "application/json",
                     "User-Agent": "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 Chrome/124.0 Safari/537.36"},
        )
        with _urlreq.urlopen(req, timeout=55) as resp:
            data = _json.loads(resp.read().decode("utf-8"))
        text = (((data.get("choices") or [{}])[0].get("message") or {}).get("content") or "").strip()
        text = re.sub(r"<think>.*?</think>", "", text, flags=re.S).strip()
        return (text, "") if text else ("", f"{model}: rỗng")
    except _urlerr.HTTPError as e:
        detail = ""
        try:
            detail = e.read().decode("utf-8")[:200]
        except Exception:
            pass
        if e.code == 429:
            return "", "RATE_LIMIT"
        if e.code in (401, 403):
            return "", "AUTH"
        return "", f"{model} -> {e.code}: {detail}"
    except Exception as e:
        return "", f"{model} -> {str(e)}"


def _openai_responses_once(api_key, model, sys_prompt, user_prompt, max_tokens, reasoning_effort,
                           json_mode=False, timeout_s=30):
    """OpenAI Responses API for higher-precision vocab extraction/classification."""
    import json as _json
    import urllib.request as _urlreq
    import urllib.error as _urlerr
    body = {
        "model": model,
        "input": [
            {"role": "system", "content": sys_prompt},
            {"role": "user", "content": user_prompt},
        ],
        "max_output_tokens": max_tokens,
        "store": False,
    }
    if json_mode:
        body["text"] = {"format": {"type": "json_object"}}
    if reasoning_effort and reasoning_effort != "none":
        body["reasoning"] = {"effort": reasoning_effort}
    try:
        req = _urlreq.Request(
            "https://api.openai.com/v1/responses",
            data=_json.dumps(body).encode("utf-8"), method="POST",
            headers={"Authorization": f"Bearer {api_key}", "Content-Type": "application/json",
                     "Accept": "application/json", "User-Agent": "IELTS-OS/1.0"},
        )
        with _urlreq.urlopen(req, timeout=timeout_s) as resp:
            data = _json.loads(resp.read().decode("utf-8"))
        text = (data.get("output_text") or "").strip()
        if not text:
            chunks = []
            for out in data.get("output") or []:
                for part in out.get("content") or []:
                    if part.get("type") in ("output_text", "text") and part.get("text"):
                        chunks.append(part.get("text"))
            text = "\n".join(chunks).strip()
        return (text, "") if text else ("", f"{model}: rỗng")
    except _urlerr.HTTPError as e:
        detail = ""
        try:
            detail = e.read().decode("utf-8")[:300]
        except Exception:
            pass
        if e.code == 429:
            return "", "RATE_LIMIT"
        if e.code in (401, 403):
            return "", "AUTH"
        return "", f"{model} -> {e.code}: {detail}"
    except Exception as e:
        return "", f"{model} -> {str(e)}"


def _gemini_chat_once(api_key, model, sys_prompt, user_prompt, max_tokens, temperature, json_mode):
    """1 lần gọi Google Gemini (định dạng riêng). Tắt 'thinking' ở 2.5 để output ổn định, đỡ rỗng."""
    import json as _json
    import urllib.request as _urlreq
    import urllib.error as _urlerr
    url = f"https://generativelanguage.googleapis.com/v1beta/models/{model}:generateContent?key={api_key}"
    gen = {"maxOutputTokens": max_tokens, "temperature": temperature}
    if json_mode:
        gen["responseMimeType"] = "application/json"
    if "2.5" in model:
        gen["thinkingConfig"] = {"thinkingBudget": 0}
    body = {
        "systemInstruction": {"parts": [{"text": sys_prompt}]},
        "contents": [{"role": "user", "parts": [{"text": user_prompt}]}],
        "generationConfig": gen,
    }
    try:
        req = _urlreq.Request(url, data=_json.dumps(body).encode("utf-8"), method="POST",
                              headers={"Content-Type": "application/json"})
        with _urlreq.urlopen(req, timeout=55) as resp:
            data = _json.loads(resp.read().decode("utf-8"))
        cands = data.get("candidates") or []
        text = ""
        if cands:
            parts = ((cands[0].get("content") or {}).get("parts") or [])
            text = "".join(p.get("text", "") for p in parts).strip()
        text = re.sub(r"<think>.*?</think>", "", text, flags=re.S).strip()
        return (text, "") if text else ("", f"gemini {model}: rỗng")
    except _urlerr.HTTPError as e:
        detail = ""
        try:
            detail = e.read().decode("utf-8")[:200]
        except Exception:
            pass
        if e.code == 429:
            return "", "RATE_LIMIT"
        if e.code in (401, 403):
            return "", "AUTH"
        return "", f"gemini {model} -> {e.code}: {detail}"
    except Exception as e:
        return "", f"gemini {model} -> {str(e)}"


def _groq_chat(sys_prompt: str, user_prompt: str, max_tokens: int = 2048, temperature: float = 0.7, json_mode: bool = False, model: str = None, reasoning_effort: str = None):
    """ĐA-PROVIDER tự xoay khi rate-limit: Groq (nhiều key) -> Cerebras -> Gemini.
    Giữ nguyên tên & signature cũ để mọi caller dùng được. Trả (text, error).
    Env: GROQ_API_KEY[_2/_3/_4], CEREBRAS_API_KEY (+CEREBRAS_MODEL), GEMINI_API_KEY (+GEMINI_MODEL)."""
    GROQ_BASE = "https://api.groq.com/openai/v1"
    CEREBRAS_BASE = "https://api.cerebras.ai/v1"
    groq_keys = [k for k in [os.environ.get("GROQ_API_KEY"), os.environ.get("GROQ_API_KEY_2"),
                             os.environ.get("GROQ_API_KEY_3"), os.environ.get("GROQ_API_KEY_4")] if k]
    cerebras_key = os.environ.get("CEREBRAS_API_KEY")
    gemini_key = os.environ.get("GEMINI_API_KEY") or os.environ.get("GOOGLE_API_KEY")
    if not (groq_keys or cerebras_key or gemini_key):
        return "", "Server chưa cấu hình GROQ_API_KEY / CEREBRAS_API_KEY / GEMINI_API_KEY."

    preferred = model or os.environ.get("GROQ_MODEL", "llama-3.3-70b-versatile")
    groq_models = [preferred] + [m for m in ["llama-3.3-70b-versatile", "llama-3.1-8b-instant"] if m != preferred]
    cerebras_models = [os.environ.get("CEREBRAS_MODEL", "gpt-oss-120b"), "llama-3.3-70b"]
    gemini_models = [os.environ.get("GEMINI_MODEL", "gemini-2.5-flash"), "gemini-2.0-flash"]

    last_err = ""
    any_rate = False

    # 1) GROQ — xoay từng key; mỗi key thử các model. Rate-limit/auth -> sang key kế.
    for key in groq_keys:
        for gm in groq_models:
            text, err = _oai_chat_once(GROQ_BASE, key, gm, sys_prompt, user_prompt, max_tokens, temperature, json_mode, reasoning_effort)
            if text:
                return text, ""
            last_err = err
            if err == "RATE_LIMIT":
                any_rate = True; break          # key này hết hạn mức -> key kế / provider kế
            if err == "AUTH":
                break

    # 2) CEREBRAS (free, OpenAI-compatible)
    if cerebras_key:
        for cm in cerebras_models:
            text, err = _oai_chat_once(CEREBRAS_BASE, cerebras_key, cm, sys_prompt, user_prompt, max_tokens, temperature, json_mode, reasoning_effort)
            if text:
                return text, ""
            last_err = err
            if err == "RATE_LIMIT":
                any_rate = True; break
            if err == "AUTH":
                break

    # 3) GEMINI (free, hạn mức rộng)
    if gemini_key:
        for gem in gemini_models:
            text, err = _gemini_chat_once(gemini_key, gem, sys_prompt, user_prompt, max_tokens, temperature, json_mode)
            if text:
                return text, ""
            last_err = err
            if err == "RATE_LIMIT":
                any_rate = True; break
            if err == "AUTH":
                break

    if any_rate:
        return "", "RATE_LIMIT"
    if last_err == "AUTH":
        return "", "API key không hợp lệ hoặc chưa kích hoạt."
    return "", last_err or "Không tạo được nội dung."


# Tương thích ngược: các nơi cũ gọi _gemini_generate -> chuyển sang Groq (bỏ tham số đa phương thức/Gemini).
def _gemini_generate(sys_prompt, user_prompt, max_tokens=2048, response_mime=None, extra_parts=None, temperature=0.7, timeout_s=45):
    return _groq_chat(sys_prompt, user_prompt, max_tokens=max_tokens, temperature=temperature, json_mode=bool(response_mime))


def _vocab_chat(sys_prompt: str, user_prompt: str, max_tokens: int = 4096, json_mode: bool = False,
                model: str = None, reasoning_effort: str = "medium"):
    """Vocab-only model router. Prefer OpenAI if configured; keep legacy free providers as fallback."""
    openai_key = os.environ.get("OPENAI_API_KEY")
    if openai_key:
        requested_openai_model = model if model and str(model).startswith("gpt-") else None
        openai_model = os.environ.get("OPENAI_VOCAB_MODEL") or os.environ.get("OPENAI_MODEL") or requested_openai_model or "gpt-5.6-terra"
        openai_effort = os.environ.get("OPENAI_VOCAB_REASONING", reasoning_effort or "medium")
        provider_timeout = max(15, min(40, int(os.environ.get("VOCAB_PROVIDER_TIMEOUT", "30") or 30)))
        text, err = _openai_responses_once(
            openai_key, openai_model, sys_prompt, user_prompt, max_tokens, openai_effort,
            json_mode=json_mode, timeout_s=provider_timeout,
        )
        if text:
            return text, ""
        # Do not burn the entire serverless invocation on a second provider after a timeout.
        # The route can return a normal JSON error and the student can retry instead.
        if "timed out" in str(err).lower() or "timeout" in str(err).lower():
            return "", "AI service timed out while generating vocabulary."
        # Bad OpenAI config should not break class. Fall back to existing providers if present.
        if err not in ("AUTH", "RATE_LIMIT"):
            print("OpenAI vocab model failed; falling back:", err)
    return _groq_chat(sys_prompt, user_prompt, max_tokens=max_tokens, temperature=0.2,
                      json_mode=json_mode, model=model, reasoning_effort=reasoning_effort)


def _friendly_err(err: str, lang: str) -> str:
    if err == "RATE_LIMIT":
        return ("Rate limit reached (too many requests in a short time). Please wait a moment and try again."
                if lang == "en" else
                "⏳ Đạt giới hạn tốc độ (quá nhiều yêu cầu trong thời gian ngắn). Vui lòng đợi chút rồi thử lại.")
    return err


@app.post("/api/ai_feedback")
async def ai_feedback(payload: Dict[str, Any] = Body(...)):
    """Sinh nhận xét cho học viên (Groq / Llama). Cần GROQ_API_KEY."""
    lang = (payload.get("lang") or "vi").lower()
    student = payload.get("studentName", "học viên")
    quiz = payload.get("quizTitle", "bài thi")
    qtype = payload.get("type", "")
    score = payload.get("score", "")
    total = payload.get("total", "")
    band = payload.get("band", "")
    weak = payload.get("weakness", "")
    wrong_count = payload.get("wrongCount", "")
    details = payload.get("details", "")

    if lang == "en":
        sys_prompt = (
            "You are an experienced IELTS examiner and coach. Based on the per-question data, write a detailed, "
            "warm, and genuinely useful feedback report for the student (and parent). Structure it as 3 short paragraphs:\n"
            "1) Open with specific encouragement tied to their band and what they did well.\n"
            "2) ANALYSE the actual mistakes: group the wrong answers, point out concrete patterns from the data "
            "(e.g. spelling/plurals in completion, confusing TRUE vs NOT GIVEN, distractors in multiple choice, "
            "paraphrase/synonym traps in matching). Reference a couple of specific wrong items as examples.\n"
            "3) Give 2-3 concrete, actionable practice steps targeting those exact weaknesses.\n"
            "Warm, professional, plain text (no markdown). About 180-260 words. Finish your final sentence."
        )
        user_prompt = (
            f"Student: {student}\nTest: {quiz} ({qtype})\nScore: {score}/{total} (Band {band}), wrong: {wrong_count}\n"
            f"Accuracy by question type: {weak}\n\nWrong answers (student vs correct):\n{details or '(none provided)'}\n\nWrite in English."
        )
    else:
        sys_prompt = (
            "Bạn là giám khảo kiêm gia sư IELTS giàu kinh nghiệm. Dựa trên dữ liệu TỪNG CÂU, viết bản nhận xét "
            "chi tiết, ấm áp, hữu ích cho học viên (và phụ huynh). Bố cục 3 đoạn ngắn:\n"
            "1) Khích lệ cụ thể, gắn với band và điểm làm tốt.\n"
            "2) PHÂN TÍCH lỗi thật: gom nhóm câu sai, chỉ quy luật cụ thể từ dữ liệu "
            "(sai chính tả/số nhiều khi điền từ; nhầm TRUE với NOT GIVEN; bẫy phương án nhiễu ở trắc nghiệm; "
            "bẫy từ đồng nghĩa/paraphrase khi nối thông tin). Dẫn vài câu sai cụ thể làm ví dụ.\n"
            "3) 2-3 bước luyện tập cụ thể, làm được ngay, nhắm đúng điểm yếu.\n"
            "Văn ấm áp, chuyên nghiệp, văn xuôi thuần (không markdown). Khoảng 180-260 từ. Viết trọn câu cuối."
        )
        user_prompt = (
            f"Học viên: {student}\nĐề: {quiz} ({qtype})\nĐiểm: {score}/{total} (Band {band}), số câu sai: {wrong_count}\n"
            f"Tỉ lệ đúng theo dạng câu: {weak}\n\nCác câu sai (HV trả lời vs đáp án đúng):\n{details or '(không có)'}\n\nViết bằng tiếng Việt."
        )

    text, err = _gemini_generate(sys_prompt, user_prompt, 2048)
    if err:
        return {"success": False, "error": _friendly_err(err, lang)}
    return {"success": True, "feedback": text}


@app.post("/api/ai_explain")
async def ai_explain(payload: Dict[str, Any] = Body(...)):
    """Giải thích ngắn vì sao 1 câu đúng/sai (dùng trong màn Review)."""
    lang = (payload.get("lang") or "vi").lower()
    question = payload.get("question", "")
    options = payload.get("options", "")
    correct = payload.get("correct", "")
    student_ans = payload.get("studentAnswer", "")
    context = (payload.get("context", "") or "")[:24000]
    reading_passage = (payload.get("readingPassage", "") or "")[:24000]
    has_ctx = bool(context.strip())
    answer_sequence = payload.get("answerSequence")
    question_index_raw = payload.get("questionIndex")
    try:
        question_index = int(question_index_raw) if question_index_raw is not None else None
    except (TypeError, ValueError):
        question_index = None
    question_type = str(payload.get("questionType", "") or "")
    question_subtype = str(payload.get("questionSubType", "") or "")
    integrated_part = int(payload.get("integratedPart", 0) or 0)
    is_vietnamese_high_school_integrated = bool(payload.get("isVietnameseHighSchoolIntegrated")) and 2 <= integrated_part <= 7
    # CHỈ yêu cầu timestamp khi transcript THẬT SỰ có mốc (m:ss - m:ss) — không có mốc mà vẫn yêu cầu là AI sẽ bịa.
    is_listening = bool(payload.get("isListening")) and bool(_TS_MARKER_RE.search(context))
    is_reading_evidence = bool(reading_passage.strip()) and not bool(payload.get("isListening"))
    ts_rule_en = (
        " TIMESTAMP RULE: the transcript contains time markers like \"(0:21 - 0:33)\" placed BEFORE each block of speech. "
        "This is an exact evidence task, not a rough location: first find the correct-answer word or phrase in the transcript, "
        "then scan UPWARD from that exact occurrence and take the NEAREST marker above it. Never choose a later related sentence, "
        "estimate, average, or use a marker after the answer-bearing words. Include exactly ONE standalone seek marker in the whole "
        "explanation, formatted only as [mm:ss] or [h:mm:ss] (e.g. marker \"(3:02 - 3:27)\" -> [3:02]); never output a time range. "
        "If the transcript has no markers, omit the timestamp entirely."
    ) if is_listening else ""
    ts_rule_vi = (
        " LUẬT MỐC THỜI GIAN: transcript có mốc dạng \"(0:21 - 0:33)\" đặt TRƯỚC mỗi khối lời thoại. "
        "Đây là việc đối chiếu CHÍNH XÁC, không phải chọn vị trí gần đúng: trước hết tìm đúng từ/cụm từ của đáp án trong transcript, "
        "sau đó dò NGƯỢC LÊN từ đúng chỗ xuất hiện đó và lấy mốc GẦN NHẤT phía trên. Tuyệt đối không lấy một câu liên quan ở phía SAU, "
        "không ước lượng, không lấy trung bình. Toàn bộ giải thích phải có đúng MỘT mốc để bấm nghe lại, chỉ ở dạng [mm:ss] hoặc "
        "[h:mm:ss] (vd mốc \"(3:02 - 3:27)\" -> [3:02]); không bao giờ ghi một khoảng thời gian. Transcript không có mốc thì bỏ hẳn timestamp."
    ) if is_listening else ""

    evidence_rule_en = (
        " READING EVIDENCE RULE: when you use evidence from the reading passage, include one to three short, continuous, "
        "word-for-word excerpts from the READING PASSAGE only. Wrap EVERY such excerpt exactly as "
        "[[EVIDENCE: exact excerpt]]. Do not place the marker around paraphrases, question wording, options, or invented text. "
        "Each marked excerpt must be the smallest useful proof and must appear exactly in the passage."
    ) if is_reading_evidence else ""
    evidence_rule_vi = (
        " LUAT DAN CHUNG READING: khi dung dan chung tu bai doc, phai dua mot den ba trich doan ngan, lien tuc va dung tung chu "
        "tu RIENG BAI DOC. Boc MOI trich doan dung cu phap [[EVIDENCE: doan trich chinh xac]]. Khong dung marker cho dien giai, "
        "loi cau hoi, lua chon hay noi dung tu bia. Moi doan trich phai la bang chung ngan nhat nhung du nghia va xuat hien nguyen van trong bai."
    ) if is_reading_evidence else ""

    integrated_rule_en = (
        " This is Integrated Part %d, modelled on Vietnam's upper-secondary national high-school English exam. "
        "Do NOT force a reading-comprehension template onto it. First identify the actual task family and teach its decision logic: "
        "for main-idea/title/inference items, test scope, central claim and degree of certainty; for gap-fill items, test grammar "
        "(word class, tense, clause structure, agreement, preposition) before meaning, then collocation and register; for dialogue/sentence "
        "ordering, establish the opening, response links, pronoun/reference chains, discourse markers and a coherent ending; for vocabulary, "
        "use the surrounding syntax and the precise sense, not a loose synonym. Explain the exam trap each wrong option represents."
    ) % integrated_part if is_vietnamese_high_school_integrated else ""
    integrated_rule_vi = (
        " Đây là Integrated Part %d, theo tư duy đề thi tiếng Anh tốt nghiệp THPT Việt Nam. "
        "KHÔNG được máy móc áp khuôn đọc đoạn văn rồi trích dẫn. Trước hết phải nhận diện đúng dạng và dạy cách ra quyết định: "
        "với ý chính/tiêu đề/suy luận, xét phạm vi, luận điểm trung tâm và mức độ khẳng định; với điền từ, xét ngữ pháp "
        "(từ loại, thì, cấu trúc mệnh đề, hòa hợp, giới từ) trước, rồi nghĩa, collocation và sắc thái; với sắp xếp hội thoại/câu, "
        "xác định câu mở, quan hệ đáp lời, đại từ thay thế, từ nối và câu kết mạch lạc; với từ vựng, dựa vào cú pháp xung quanh "
        "và nghĩa chính xác, không chọn từ đồng nghĩa chung chung. Chỉ rõ bẫy mà từng phương án sai đang tạo ra."
    ) % integrated_part if is_vietnamese_high_school_integrated else ""

    if lang == "en":
        sys_prompt = (
            "You are an expert English-test tutor. Give a clear, precise and genuinely teachable explanation for one reviewed item. "
            "First identify what the question is testing; do not pretend every item is a locate-and-quote reading question. "
            "Use this teaching sequence: explain the key decision, prove why the correct answer fits, then account for EVERY remaining "
            "option when options are provided. For each wrong option, name its specific flaw: wrong grammar/collocation, wrong reference, "
            "too broad/narrow/absolute, only a true detail rather than the main point, reversed logic, unsupported inference, or incoherent "
            "discourse order. For a gap without options, state the required grammar and lexical pattern and why the student's form fails. "
            "Quote the source only when it is relevant and actually supports the point; never invent a quotation or claim that an answer must "
            "appear word-for-word in the source. If the student was correct, still explain why the other options lose. End with one short, "
            "reusable solving habit. Write 5-9 concise but detailed sentences in plain text, with short labels if helpful; no markdown." + integrated_rule_en + ts_rule_en + evidence_rule_en
        )
        user_prompt = (
            f"SOURCE TEXT (read it fully):\n{context if has_ctx else '(no source text was provided for this item)'}\n\n"
            f"Question type: {question_type or '(unspecified)'}; subtype: {question_subtype or '(none)'}\n"
            f"Question: {question}\nOptions: {options or '(n/a)'}\n"
            f"Correct answer: {correct}\nStudent answered: {student_ans or '(blank)'}\n\nExplain in English."
        )
    else:
        sys_prompt = (
            "Bạn là gia sư luyện thi tiếng Anh giàu kinh nghiệm. Hãy giải thích một câu trong phần Review theo cách dễ hiểu, chính xác "
            "và thực sự giúp học viên tự làm được lần sau. Trước tiên nhận diện câu hỏi đang kiểm tra gì; không được giả vờ mọi dạng đều "
            "là tìm câu trong bài rồi trích dẫn. Đi theo trình tự dạy học: nêu điểm mấu chốt để ra quyết định, chứng minh vì sao đáp án đúng "
            "phù hợp, rồi phân tích TẤT CẢ phương án còn lại nếu có lựa chọn. Mỗi phương án sai phải nêu lỗi RIÊNG: sai ngữ pháp/collocation, "
            "sai đối tượng tham chiếu, quá rộng/quá hẹp/quá tuyệt đối, chỉ đúng một chi tiết nhưng không phải ý chính, đảo quan hệ nguyên nhân-kết quả, "
            "suy luận không được hỗ trợ, hoặc phá vỡ mạch hội thoại/văn bản. Với câu điền không có lựa chọn, nói rõ cấu trúc ngữ pháp và mẫu từ vựng "
            "cần dùng, cũng như vì sao cách điền của học viên chưa đúng. Chỉ trích dẫn văn bản khi trích dẫn thực sự liên quan và chứng minh được ý; "
            "tuyệt đối không bịa trích dẫn hoặc ép đáp án phải xuất hiện nguyên văn. Học viên làm đúng vẫn cần biết vì sao các đáp án khác bị loại. "
            "Kết thúc bằng một mẹo làm bài ngắn có thể áp dụng lại. Viết 5-9 câu gọn nhưng chi tiết, văn xuôi thuần, có thể dùng nhãn ngắn; không markdown." + integrated_rule_vi + ts_rule_vi + evidence_rule_vi
        )
        user_prompt = (
            f"VĂN BẢN NGUỒN (đọc hết):\n{context if has_ctx else '(không có văn bản nguồn cho câu này)'}\n\n"
            f"Dạng câu: {question_type or '(chưa xác định)'}; dạng phụ: {question_subtype or '(không có)'}\n"
            f"Câu hỏi: {question}\nLựa chọn: {options or '(không có)'}\n"
            f"Đáp án đúng: {correct}\nHọc viên chọn: {student_ans or '(bỏ trống)'}\n\nGiải thích bằng tiếng Việt."
        )

    text, err = _gemini_generate(sys_prompt, user_prompt, 1200)
    if err:
        return {"success": False, "error": _friendly_err(err, lang)}
    # HẬU KIỂM: xóa mọi [mm:ss] không khớp mốc thật trong transcript (chống AI bịa/ước lượng).
    text = _filter_fake_timestamps(
        text,
        context,
        correct if is_listening else "",
        lang,
        answer_sequence if is_listening else None,
        question_index if is_listening else None,
    )
    if is_reading_evidence:
        def _norm_evidence(value: str) -> str:
            return re.sub(r"\s+", " ", str(value or "").replace("’", "'").replace("“", '"').replace("”", '"')).strip().lower()

        passage_norm = _norm_evidence(reading_passage)

        def _keep_real_evidence(match: re.Match) -> str:
            quote = match.group(1).strip()
            quote_norm = _norm_evidence(quote)
            if 4 <= len(quote_norm) <= 300 and quote_norm in passage_norm:
                return f"[[EVIDENCE: {quote}]]"
            return f'"{quote}"'

        text = re.sub(r"\[\[EVIDENCE:\s*([\s\S]*?)\s*\]\]", _keep_real_evidence, text, flags=re.IGNORECASE)
    return {"success": True, "explanation": text}


def _parse_ai_items(text: str):
    """Bóc danh sách item từ output AI: chấp nhận {items:[...]}, mảng thuần, hoặc [ ... ] lẫn trong văn bản."""
    import json as _json
    raw = (text or "").strip()
    if raw.startswith("```"):
        raw = raw.strip("`")
        if raw.lower().startswith("json"):
            raw = raw[4:]
    try:
        parsed = _json.loads(raw)
        if isinstance(parsed, list):
            return parsed
        if isinstance(parsed, dict):
            arr = parsed.get("items") or next((v for v in parsed.values() if isinstance(v, list)), None)
            if isinstance(arr, list):
                return arr
    except Exception:
        pass
    s, e = raw.find("["), raw.rfind("]")
    if s != -1 and e != -1 and e > s:
        try:
            arr = _json.loads(raw[s:e + 1])
            if isinstance(arr, list):
                return arr
        except Exception:
            pass
    return None


# GIAI ĐOẠN B — cây quyết định + few-shot cho model reasoning phân loại CHÍNH XÁC.
_VOCAB_CLASSIFY_SYS = (
    "You are an expert IELTS lexicographer. Analyze EACH candidate into observable lexical FEATURES. "
    "Do not jump straight to a category. Reason silently, output ONLY the JSON.\n"
    "STEP 1 — THE KEEP TEST. Set keep=false (DROP) if the item is ANY of:\n"
    "- an AD-HOC word combination that is NOT an established expression found in a dictionary/coursebook — just words that "
    "happen to sit next to each other, e.g. 'carved deep', 'retreating glaciers', 'kettle ponds', 'low-lying areas'.\n"
    "- a NOVEL or LITERARY metaphor coined in the passage (NOT a real, widely-used idiom), e.g. 'formative hothouse'.\n"
    "- a bare verb conjugated in a tense, e.g. 'have been grown', 'are generally cultivated', 'have made', 'was developed'.\n"
    "- a sentence fragment, a proper noun, or a trivial A1/A2 word ('people', 'because', 'important').\n"
    "RULE: if a combo is ad-hoc but contains a valuable single word, DROP the combo (the single word is handled separately). "
    "When unsure whether a multi-word phrase is truly established, prefer keep=false. Precision over quantity.\n"
    "STEP 2 — CATEGORY for KEPT items, by DECISION TREE (stop at first match):\n"
    "1) SINGLE word, no spaces => \"word\".\n"
    "2) ABSTRACT grammatical template with a fillable slot — 'not only ... but also', 'the more ... the more', cleft "
    "'It was X that ...', inversion 'Never had he ...', reporting 'is said to ...', conditionals => \"grammar\".\n"
    "3) BASE VERB + a TRUE ADVERBIAL particle (up/out/down/off/away/back/over/through/along/around/apart/aside; plus on/in/into/onto "
    "WHEN adverbial) forming an ESTABLISHED phrasal verb where the verb's meaning is changed/completed idiomatically OR the particle "
    "can move after a noun object ('carry out the plan'='carry the plan out') — 'carry out', 'put up with', 'set off', 'point out', "
    "'give up', 'bring about', 'go into', 'look into', 'turn into' => \"phrasal_verb\".\n"
    "   !! NOT a phrasal verb: a PREPOSITIONAL VERB = verb keeps its literal meaning + a DEPENDENT PREPOSITION that cannot move "
    "('depend on', 'rely on', 'focus on', 'insist on', 'consist of', 'approve of', 'refer to', 'adhere to', 'contribute to', "
    "'react to', 'cope with', 'comply with', 'differ from', 'refrain from', 'account for', 'believe in', 'result in', 'embark on', "
    "'delve into', 'preside over', 'worry about', 'guard against') => these are \"collocation\", NOT phrasal_verb. "
    "Also 'look at'/'listen to'/'wait for' (basic verb + prep) are worthless => keep:false.\n"
    "4) ESTABLISHED idiom — a FIXED, dictionary-attested FIGURATIVE expression ('a blessing in disguise', 'break new ground', "
    "'the tip of the iceberg'). A metaphor invented in the text is NOT an idiom. "
    "HARD RULE: most verb-led phrases that END IN A PREPOSITION/PARTICLE ('look forward to', 'get to', 'make a note of', "
    "'come up with') are NOT idioms — they are phrasal_verb or collocation. EXCEPTION: dictionary-attested figurative idioms "
    "such as 'keep an eye on' and 'turn a blind eye to' remain idiom.\n"
    "5) OTHERWISE a multi-word LITERAL collocation / linking phrase / prepositional verb ('in terms of', 'as a result of', "
    "'pose a threat', 'play a vital role', 'heavily reliant on', 'depend on', 'consist of') => \"collocation\".\n"
    "WORKED EXAMPLES (word -> category, keep): meticulous->word,true | formative->word,true | hothouse->word,true | "
    "carve->word,true | carry out->phrasal_verb,true | put up with->phrasal_verb,true | go into->phrasal_verb,true | "
    "look into->phrasal_verb,true | give up->phrasal_verb,true | look at->(any),false | refer to->collocation,true | "
    "depend on->collocation,true | rely on->collocation,true | consist of->collocation,true | focus on->collocation,true | "
    "in terms of->collocation,true | pose a threat->collocation,true | break new ground->idiom,true | "
    "not only ... but also->grammar,true | carved deep->(any),false | formative hothouse->(any),false | "
    "retreating glaciers->(any),false | have been grown->(any),false | are generally cultivated->(any),false\n"
    "OUTPUT: a JSON array, one object per input item IN THE SAME ORDER. I parse it programmatically. "
    "Use EXACTLY these keys: word, keep, established, lexical_type, canonical_form, figurative, confidence. "
    "lexical_type MUST be one of single_word|grammar_template|phrasal_verb|prepositional_verb|collocation|idiom|not_lexical. "
    "canonical_form is the dictionary head form (e.g. input 'burn off excess energy' => 'burn off'; "
    "input 'you're in luck' => 'be in luck'). figurative is true ONLY when the WHOLE expression has a conventional "
    "non-literal meaning. confidence MUST be high|medium|low. Example: "
    "[{\"word\":\"up to date\",\"keep\":true,\"established\":true,\"lexical_type\":\"collocation\","
    "\"canonical_form\":\"up to date\",\"figurative\":false,\"confidence\":\"high\"}]. "
    "Output the JSON array and nothing after it."
)


def _is_abstract_grammar(word: str) -> bool:
    """True nếu 'word' là TEMPLATE ngữ pháp trừu tượng (có slot), False nếu chỉ là instance chia thì.
    Chặn rác kiểu 'have been grown', 'must act', 'is said to have' lọt vào nhóm grammar."""
    w = str(word or "").strip()
    if not w:
        return False
    # dấu hiệu template: ellipsis, dấu +, gạch chéo lựa chọn
    if any(m in w for m in ("...", "…", " + ", "+", "/")):
        return True
    # placeholder ký hiệu: S / V / O / N / sb / sth / sth's / adj / inf
    if re.search(r"\b(S|V|O|N|sb|sth|adj|inf|to-infinitive|infinitive)\b", w):
        return True
    # tên/loại cấu trúc được gọi tên rõ ràng
    if re.search(r"(inversion|passive|cleft|conditional|comparative|correlative|relative clause|"
                 r"reported speech|subjunctive|gerund|đảo ngữ|bị động|điều kiện|so sánh)", w, re.I):
        return True
    return False


# Bộ dò ngữ pháp TẤT ĐỊNH: quét material bằng regex, sinh grammar item dạng pattern chuẩn.
# Mỗi entry: (key, regex, pattern_word_en, pos, meaning_vi, meaning_en)
_GRAMMAR_PATTERNS = [
    ("not_only",
     re.compile(r"\bnot only\b[^.?!]*\bbut\b[^.?!]*\balso\b", re.I),
     "Not only + auxiliary + S ..., but S also ... (inversion)",
     "đảo ngữ nhấn mạnh",
     "Đảo ngữ với 'Not only ... but ... also' để nhấn mạnh hai vế. Sau 'Not only' đảo trợ động từ lên trước chủ ngữ.",
     "Inversion with 'Not only ... but ... also' for emphasis; the auxiliary comes before the subject after 'Not only'."),
    ("neg_inversion",
     re.compile(r"\b(never|rarely|seldom|hardly|scarcely|no sooner|nowhere|little)\s+(had|has|have|did|does|was|were|is|are|will|would|can|could)\s+\w+", re.I),
     "Negative adverbial + auxiliary + S + V (inversion)",
     "đảo ngữ phủ định",
     "Đảo ngữ sau trạng từ phủ định đầu câu (Never/Rarely/Seldom/Hardly...): đưa trợ động từ lên trước chủ ngữ.",
     "Inversion after a fronted negative adverbial (Never/Rarely/Seldom/Hardly...): the auxiliary precedes the subject."),
    ("passive_report",
     re.compile(r"\b(is|are|was|were|has been|have been)\s+(said|believed|thought|reported|known|expected|considered|claimed|estimated|alleged|presumed|rumou?red)\s+to\b", re.I),
     "S + be + said/believed/thought + to-infinitive (passive reporting)",
     "bị động tường thuật",
     "Bị động tường thuật: 'S + be + said/believed/thought + to + V' để truyền đạt thông tin một cách khách quan.",
     "Passive reporting structure 'S + be + said/believed/thought + to-infinitive' to report information impersonally."),
    ("cleft_it",
     re.compile(r"\bit\s+(is|was)\b[\w\s,'-]{1,45}?\b(that|who|which)\b", re.I),
     "It + be + emphasised part + that/who ... (cleft sentence)",
     "câu chẻ nhấn mạnh",
     "Câu chẻ (cleft) 'It is/was ... that/who ...' để nhấn mạnh một thành phần trong câu.",
     "Cleft sentence 'It is/was ... that/who ...' used to emphasise one element of the clause."),
    ("correlative_comp",
     re.compile(r"\bthe\s+\w+(?:er|more|less|better|worse)\b[^.?!]*\bthe\s+\w+(?:er|more|less|better|worse)\b", re.I),
     "The + comparative ..., the + comparative ... (correlative)",
     "so sánh tương quan",
     "Cấu trúc 'The + so sánh hơn ..., the + so sánh hơn ...' (càng ... càng ...).",
     "Correlative comparative 'The + comparative ..., the + comparative ...' (the more ... the more ...)."),
    ("third_cond",
     re.compile(r"\bif\b[^.?!]*\bhad\b[^.?!]*\bwould have\b", re.I),
     "If + S + had + V3, S + would have + V3 (third conditional)",
     "câu điều kiện loại 3",
     "Câu điều kiện loại 3: 'If + S + had + V3, S + would have + V3' nói về điều trái với quá khứ.",
     "Third conditional 'If + had + past participle, ... would have + past participle' for unreal past situations."),
]


def _detect_grammar(source: str, lang="vi", max_items: int = 4):
    """Quét material, trả về các grammar item dạng pattern chuẩn (đã grounded sẵn bằng câu khớp)."""
    if not source:
        return []
    meaning_key = "meaning_vi" if lang == "vi" else "meaning_en"
    # tách câu thô để gắn source_sentence
    sentences = re.split(r"(?<=[.?!])\s+", source)
    out = []
    for key, rx, pat_word, pos, mvi, men in _GRAMMAR_PATTERNS:
        m = rx.search(source)
        if not m:
            continue
        seg = m.group(0)
        sent = next((s.strip() for s in sentences if seg[:25].lower() in s.lower()), seg.strip())
        item = {
            "word": pat_word, "category": "grammar", "cefr": "B2", "pos": pos,
            "phonetic": "", "source_sentence": sent[:300], "example": sent[:300],
            meaning_key: (mvi if lang == "vi" else men), "_grammar_key": key,
        }
        out.append(item)
        if len(out) >= max_items:
            break
    return out


def _classify_vocab_items(items, lang="vi"):
    """GIAI ĐOẠN B: phân loại lại từng candidate bằng model reasoning của Groq.
    An toàn: lỗi/parse-fail -> trả nguyên items (giữ category thô từ giai đoạn A)."""
    cand = [it for it in items if isinstance(it, dict) and str(it.get("word", "")).strip()]
    if not cand:
        return items
    lines = []
    for i, it in enumerate(cand, 1):
        ctx = re.sub(r"\s+", " ", str(it.get("source_sentence") or it.get("example") or "")).strip()[:180]
        lines.append(f'{i}. "{str(it.get("word")).strip()}" — context: {ctx or "(none)"}')
    user_p = (
        "Classify EVERY item below. Use the context only to judge whether a multi-word phrase is figurative (idiom) "
        "or literal (collocation). Items:\n" + "\n".join(lines) +
        "\n\nReturn the JSON array now, one object per item, in the SAME order."
    )
    model = os.environ.get("VOCAB_CLASSIFY_MODEL") or os.environ.get("GROQ_REASON_MODEL", "openai/gpt-oss-120b")
    text, err = _vocab_chat(_VOCAB_CLASSIFY_SYS, user_p, max_tokens=3072, json_mode=False,
                            model=model, reasoning_effort=os.environ.get("VOCAB_CLASSIFY_REASONING", "medium"))
    if err or not text:
        return items
    arr = _parse_ai_items(text)
    if not isinstance(arr, list) or not arr:
        return items
    valid_types = {
        "single_word", "grammar_template", "phrasal_verb", "prepositional_verb",
        "collocation", "idiom", "not_lexical",
    }

    def _norm_bool(value, default=None):
        if isinstance(value, bool):
            return value
        normalized = str(value or "").strip().lower()
        if normalized in {"true", "1", "yes"}:
            return True
        if normalized in {"false", "0", "no"}:
            return False
        return default

    def _norm_type(value):
        normalized = str(value or "").strip().lower().replace("-", "_").replace(" ", "_")
        aliases = {
            "word": "single_word", "grammar": "grammar_template", "fixed_phrase": "collocation",
            "prepositional_phrase": "collocation", "expression": "collocation", "drop": "not_lexical",
        }
        normalized = aliases.get(normalized, normalized)
        return normalized if normalized in valid_types else None

    def _norm_form(value):
        return re.sub(r"\s+", " ", str(value or "").strip().lower())

    def _form_is_grounded(form, item):
        if not form or len(form.split()) < 2 or len(form.split()) > 5:
            return False
        context = _norm_form(item.get("source_sentence") or item.get("example") or "")
        context = re.sub(r"[^a-z0-9' ]", " ", context)
        context = re.sub(r"\s+", " ", context).strip()
        return bool(context and re.search(rf"(?<![a-z0-9]){re.escape(form)}(?![a-z0-9])", context))

    by_word = {}
    for o in arr:
        if isinstance(o, dict):
            w = re.sub(r"\s+", " ", str(o.get("word", "")).strip().lower())
            if w:
                by_word[w] = o
    ordered = [o for o in arr if isinstance(o, dict)]
    use_order = (len(ordered) == len(cand))
    for idx, it in enumerate(cand):
        w = re.sub(r"\s+", " ", str(it.get("word", "")).strip().lower())
        entry = by_word.get(w)
        if entry is None and use_order:
            entry = ordered[idx]
        if entry is None:
            continue
        keep = _norm_bool(entry.get("keep"), True)
        established = _norm_bool(entry.get("established"), None)
        figurative = _norm_bool(entry.get("figurative"), None)
        lexical_type = _norm_type(entry.get("lexical_type") or entry.get("category"))
        confidence = str(entry.get("confidence") or "medium").strip().lower()
        if confidence not in {"high", "medium", "low"}:
            confidence = "medium"

        if not keep or established is False or lexical_type == "not_lexical":
            it["category"] = "__drop__"
            continue

        category_by_type = {
            "single_word": "word", "grammar_template": "grammar", "phrasal_verb": "phrasal_verb",
            "prepositional_verb": "collocation", "collocation": "collocation", "idiom": "idiom",
        }
        category = category_by_type.get(lexical_type)
        if category:
            # An idiom requires positive figurative evidence. Literal fixed phrases are collocations.
            if category == "idiom" and figurative is not True:
                category = "collocation"
            # Low-confidence particle analysis falls back to the broader literal bucket.
            if category == "phrasal_verb" and confidence == "low":
                category = "collocation"
            it["category"] = category

        canonical = _norm_form(entry.get("canonical_form"))
        if lexical_type == "phrasal_verb" and _form_is_grounded(canonical, it):
            it["word"] = canonical
        it["_classifier_figurative"] = figurative
        it["_classifier_established"] = established
        it["_classifier_confidence"] = confidence
    return items


# ===== Hằng số phân loại (module-scope để cả pass chính & pass bù dùng chung) =====
_VALID_CATS = {"word", "phrasal_verb", "idiom", "collocation", "grammar"}
_CAT_ALIAS = {
    "phrasal verb": "phrasal_verb", "phrasalverb": "phrasal_verb", "phrasal": "phrasal_verb",
    "idioms": "idiom",
    "expression": "collocation", "fixed expression": "collocation", "fixed_expression": "collocation",
    "collocations": "collocation", "prepositional phrase": "collocation", "linking phrase": "collocation",
    "grammar structure": "grammar", "structure": "grammar",
    "single word": "word", "vocabulary": "word", "vocab": "word",
}
_PREP_NOT_PARTICLE = {"to", "at", "for", "from", "of", "with", "against", "upon"}
# 'into'/'onto' LÀ particle phrasal verb thật ('go into', 'look into', 'run into', 'turn into', 'bump into').
_PARTICLES = {"up", "out", "down", "off", "on", "in", "into", "onto", "away", "back", "over", "through",
              "along", "around", "about", "apart", "aside", "forward", "forth", "across",
              "by", "past", "under", "together", "ahead", "round"}
_TRANSPARENT_CONNECTORS = {
    "in response to", "in terms of", "as a result of", "with regard to", "with respect to",
    "in addition to", "in contrast to", "on the basis of", "in relation to", "in line with",
    "in accordance with", "by means of", "in favour of", "in favor of", "in the context of",
    "as opposed to", "due to", "owing to", "prior to", "regardless of", "apart from",
    "in comparison with", "in comparison to", "with the aim of", "for the purpose of",
    "on behalf of", "in the absence of", "in the event of", "by virtue of",
}
_TRUE_PHRASAL_VERBS = {
    "break down", "bring about", "burn off", "carry out", "come across", "come up with",
    "cut down on", "find out", "get away with", "give up", "go into", "go through", "look into",
    "look forward to", "make up for", "point out", "put off", "put up with", "run into", "set off",
    "set up", "take on", "take over", "turn into", "turn out", "work out", "stumble upon",
}
_IDIOM_WHITELIST = {
    "a blessing in disguise", "a double-edged sword", "a drop in the ocean", "a piece of cake",
    "be in luck", "break new ground", "bear fruit", "in luck", "keep an eye on", "the tip of the iceberg",
    "turn a blind eye to", "under the weather", "when it comes to", "you are in luck", "you're in luck",
}
_LITERAL_FIXED_PHRASES = {
    "come close to doing", "make a note of", "up to date",
}
# Giới từ TRƠN (gần như không bao giờ là particle phrasal verb).
_TRIVIAL_PREPS = {"at", "to", "for", "of", "from"}
# Động từ cơ bản A1/A2: 'verb + giới từ trơn' (vd 'look at', 'listen to', 'wait for') = vô giá trị flashcard -> loại.
_BASIC_VERBS = {
    "look", "listen", "wait", "talk", "speak", "think", "go", "come", "get", "put", "take",
    "make", "see", "give", "find", "want", "like", "need", "know", "say", "tell", "ask",
    "work", "play", "live", "feel", "leave", "keep", "let", "run", "move", "turn", "start",
    "stop", "help", "try", "call", "use", "show", "walk", "sit", "stand", "hear", "watch",
    "read", "write", "buy", "sell", "eat", "drink", "open", "close", "send", "meet", "pay",
    "hope", "mean", "hold", "add", "point",
}
# PREPOSITIONAL VERBS = động từ + GIỚI TỪ phụ thuộc (verb GIỮ NGHĨA gốc), KHÔNG phải phrasal verb.
# Bẫy kinh điển: giới từ 'on/in/into/over/about/upon' nằm trong _PARTICLES nên 'depend on', 'rely on'... bị
# gán nhầm phrasal_verb. Whitelist này KHOÁ chúng về collocation, chặn các luật particle bên dưới lật ngược.
# CHỈ liệt kê những combo giới-từ RÕ RÀNG (verb không đổi nghĩa); các combo idiomatic thật ('bring about',
# 'go through', 'take over', 'count on', 'catch on', 'carry on'...) CỐ Ý để ngoài -> vẫn là phrasal verb.
_PREP_VERBS = {
    # + on
    "depend on", "rely on", "focus on", "insist on", "concentrate on", "embark on", "dwell on",
    "capitalize on", "border on", "prey on", "hinge on", "thrive on", "feed on", "impose on", "reflect on",
    # + in
    "believe in", "result in", "specialize in", "invest in", "engage in", "participate in", "persist in",
    "confide in", "indulge in", "revel in", "partake in", "culminate in",
    # + into
    "delve into", "tap into", "inquire into",
    # + over
    "preside over", "brood over", "mull over", "pore over", "gloss over",
    # + about
    "worry about", "complain about", "fret about", "quibble about",
    # + upon
    "rely upon", "embark upon", "touch upon", "reflect upon", "frown upon", "stumble upon", "depend upon",
    "prey upon", "seize upon", "dwell upon", "hinge upon",
    # + to (giới từ trơn — bổ sung để chắc chắn ra collocation, không phải phrasal)
    "adhere to", "conform to", "aspire to", "allude to", "succumb to", "object to", "resort to",
    "amount to", "subscribe to", "contribute to", "react to", "respond to", "relate to", "revert to",
    "pertain to", "testify to", "attest to", "cater to", "refer to",
    # + from
    "refrain from", "abstain from", "deviate from", "stem from", "derive from", "benefit from",
    "suffer from", "recover from", "differ from",
    # + with
    "cope with", "comply with", "interfere with", "sympathize with", "cooperate with", "collide with",
    "tamper with", "dispense with", "grapple with", "reckon with", "meddle with",
    # + for
    "account for", "cater for", "provide for", "compete for", "yearn for", "long for", "atone for",
    "vouch for", "clamour for", "clamor for", "opt for",
    # + of
    "consist of", "approve of", "dispose of", "conceive of", "boast of", "despair of",
    # + at
    "hint at", "marvel at", "excel at", "balk at", "gaze at",
    # + against
    "protest against", "guard against", "discriminate against", "rebel against", "militate against",
}


def _normalize_vocab_items(items):
    """Chuẩn hóa + siết phân loại 5 loại (precision cao). Đánh dấu '__drop__' item rác.
    Tách ra hàm để cả pass chính lẫn pass bù phủ-loại đều dùng được."""
    for it in items:
        if not isinstance(it, dict):
            continue
        c = str(it.get("category", "") or "").strip().lower().replace("-", "_")
        if c != "__drop__":
            c = c if c in _VALID_CATS else _CAT_ALIAS.get(c.replace("_", " "), _CAT_ALIAS.get(c, "word"))
        w = re.sub(r"\s+", " ", str(it.get("word", "") or "").strip().lower())
        classifier_figurative = it.pop("_classifier_figurative", None)
        classifier_established = it.pop("_classifier_established", None)
        classifier_confidence = str(it.pop("_classifier_confidence", "") or "").lower()
        if c == "__drop__":
            it["category"] = c
            continue
        # Canonicalise a phrasal verb that the extraction pass returned together with its object.
        phrasal_head = next((pv for pv in sorted(_TRUE_PHRASAL_VERBS, key=len, reverse=True)
                             if w == pv or w.startswith(pv + " ")), None)
        if phrasal_head:
            c = "phrasal_verb"
            if w != phrasal_head:
                w = phrasal_head
                it["word"] = phrasal_head
        # KHOÁ prepositional verb -> collocation TRƯỚC mọi luật particle (nếu không 'depend on' sẽ bị lật thành phrasal).
        _prep_locked = (w in _PREP_VERBS or (w.startswith("to ") and w[3:] in _PREP_VERBS)) and w not in _TRUE_PHRASAL_VERBS
        if _prep_locked:
            c = "collocation"
        elif w in _LITERAL_FIXED_PHRASES:
            c = "collocation"
        elif w in _IDIOM_WHITELIST:
            c = "idiom"
        elif w in _TRUE_PHRASAL_VERBS:
            c = "phrasal_verb"
        elif w in _TRANSPARENT_CONNECTORS:
            c = "collocation"
        elif w and " " not in w and c in ("idiom", "collocation"):
            c = "word"
        elif c == "grammar" and w and len(w.split()) <= 4 and re.match(r"^(in|on|at|by|with|for|as|of|to|from|due|owing|prior|apart|regardless)\b", w):
            c = "collocation"
        # Multi-word bị gán nhầm 'word' -> định tuyến lại theo hình thái.
        if c == "word" and " " in w:
            _mt = w.split()
            if len(_mt) == 2 and _mt[1] in _PARTICLES:
                c = "phrasal_verb"
            elif len(_mt) == 3 and _mt[1] in _PARTICLES and _mt[2] in ("with", "to", "for", "on"):
                c = "phrasal_verb"
            else:
                c = "collocation"
        # ====== IDIOM GATE (chống LLM gán 'idiom' bừa) ======
        # Idiom THẬT là cụm HÌNH TƯỢNG, cố định, thường KẾT THÚC BẰNG DANH TỪ/nội dung
        # ('break new ground', 'the tip of the iceberg', 'a blessing in disguise', 'bear fruit').
        # Ngược lại, cụm động từ KẾT THÚC BẰNG GIỚI TỪ/PARTICLE ('look forward to', 'get to',
        # 'make a note of', 'keep an eye on', 'come up with') KHÔNG phải idiom -> hạ về collocation,
        # rồi để tầng collocation/phrasal bên dưới tinh chỉnh tiếp (thành phrasal_verb hoặc drop nếu rác).
        if c == "idiom":
            _idt = w.split()
            _last = _idt[-1] if _idt else ""
            _first = _idt[0] if _idt else ""
            if w in _IDIOM_WHITELIST:
                c = "idiom"
            elif (classifier_figurative is True and classifier_established is not False
                  and classifier_confidence != "low"):
                c = "idiom"
            elif len(_idt) == 2 and _first not in ("the", "a", "an", "in", "on", "at", "by") and _last in _PARTICLES:
                c = "phrasal_verb"
            elif (len(_idt) >= 2
                  and (_last in _PARTICLES or _last in _PREP_NOT_PARTICLE or _last in _TRIVIAL_PREPS)):
                # Kết thúc bằng giới từ/particle => KHÔNG phải idiom.
                c = "collocation"
            else:
                # If the specialist pass is unavailable, never promote an unverified literal phrase to idiom.
                c = "collocation"
        # 'look at', 'listen to', 'wait for'... = prepositional verb cơ bản, KHÔNG phải collocation -> loại.
        # _prep_locked (depend on, rely on...) = prepositional verb hợp lệ -> GIỮ collocation, KHÔNG lật thành phrasal.
        if c == "collocation" and not _prep_locked:
            _ct = w.split()
            if len(_ct) == 2 and _ct[1] in _TRIVIAL_PREPS and _ct[0] in _BASIC_VERBS:
                c = "__drop__"
            # 2 từ kết thúc bằng particle thật ('go into', 'look into', 'carry out') = phrasal verb bị gán nhầm collocation.
            elif len(_ct) == 2 and _ct[1] in _PARTICLES:
                c = "phrasal_verb"
            # Phrasal-prepositional 3 từ ('come up with', 'look forward to', 'get away with') gán nhầm collocation.
            elif len(_ct) == 3 and _ct[1] in _PARTICLES and _ct[2] in ("with", "to", "for", "on", "of", "from"):
                c = "phrasal_verb"
        if c == "phrasal_verb":
            _wt = w.split()
            # Bỏ 'to ' đầu ('to carry out' -> 'carry out')
            if _wt and _wt[0] == "to" and len(_wt) >= 3:
                _wt = _wt[1:]; w = " ".join(_wt); it["word"] = w
            if w in _TRUE_PHRASAL_VERBS:
                c = "phrasal_verb"
            elif (_wt and _wt[0] in ("being", "been", "be", "is", "are", "was", "were")) or len(_wt) < 2:
                c = "__drop__"
            elif not any(p in _PARTICLES for p in _wt[1:]):
                # 'deal with', 'adhere to', 'rely upon'... = prepositional verb: KHÔNG vứt, chuyển collocation.
                # Riêng động từ cơ bản + giới từ trơn ('look at') vẫn loại vì vô giá trị flashcard.
                if _wt[0] in _BASIC_VERBS and _wt[1] in _TRIVIAL_PREPS:
                    c = "__drop__"
                else:
                    c = "collocation"
        _mean = str(it.get("meaning_en") or it.get("meaning_vi") or it.get("meaning") or "").lower()
        _w0 = w.split()[0] if w.split() else ""
        if (c in ("collocation", "word", "idiom")
                and _w0 in ("have", "has", "had", "is", "are", "was", "were", "am", "be", "been", "being", "will", "would")
                and (" tense" in _mean or "passive voice" in _mean or "passive tense" in _mean
                     or "thì " in _mean or "bị động" in _mean or "thể bị động" in _mean)):
            c = "__drop__"
        if c == "grammar" and not _is_abstract_grammar(it.get("word", "")):
            c = "__drop__"
        it["category"] = c
    return [it for it in items if not (isinstance(it, dict) and it.get("category") == "__drop__")]


def _ground_vocab_items(items, source):
    """Chống bịa: CHÍNH cụm từ (hoặc trích dẫn grammar) phải nằm trong material. Trả (verified, dropped)."""
    def _norm(s: str) -> str:
        return re.sub(r"\s+", " ", re.sub(r"[^a-z0-9 ]", " ", str(s or "").lower())).strip()
    src_norm = _norm(source)
    if not src_norm:
        return items, 0
    src_tokens = src_norm.split()

    def _tok_match(a: str, b: str) -> bool:
        if a == b:
            return True
        if len(a) >= 4 and len(b) >= 4:
            if b.startswith(a[:max(4, len(a) - 2)]) or a.startswith(b[:max(4, len(b) - 2)]):
                return True
        return False

    def _phrase_in_source(wt) -> bool:
        n = len(wt)
        if n == 0:
            return False
        if n == 1:
            return any(_tok_match(wt[0], st) for st in src_tokens)
        for i in range(len(src_tokens) - n + 1):
            if all(_tok_match(wt[j], src_tokens[i + j]) for j in range(n)):
                return True
        return False

    def _loose_seq_in_source(wt) -> bool:
        # Phrasal verb TÁCH RỜI: 'carry it out', 'set them off', 'put the plan into action'.
        # Khớp động từ rồi các particle theo THỨ TỰ, cho xen tối đa ~4 từ giữa các phần.
        n = len(wt)
        if n < 2:
            return _phrase_in_source(wt)
        for i in range(len(src_tokens)):
            if not _tok_match(wt[0], src_tokens[i]):
                continue
            pos, ok = i, True
            for tok in wt[1:]:
                found = -1
                for j in range(pos + 1, min(len(src_tokens), pos + 6)):
                    if _tok_match(tok, src_tokens[j]):
                        found = j
                        break
                if found == -1:
                    ok = False
                    break
                pos = found
            if ok:
                return True
        return False

    def _quote_in_source(ss: str) -> bool:
        toks = ss.split()
        return len(toks) >= 5 and (ss in src_norm or " ".join(toks[:5]) in src_norm or " ".join(toks[-5:]) in src_norm)

    def _grounded(it: dict) -> bool:
        if not isinstance(it, dict):
            return False
        cat = it.get("category", "word")
        if cat == "grammar":
            return _quote_in_source(_norm(it.get("source_sentence", "")))
        wt = _norm(it.get("word", "")).split()
        if cat == "phrasal_verb":
            return _loose_seq_in_source(wt)   # cho phép tách rời
        return _phrase_in_source(wt)

    verified = [it for it in items if _grounded(it)]
    return verified, len(items) - len(verified)


@app.post("/api/ai_vocab")
async def ai_vocab(payload: Dict[str, Any] = Body(...)):
    """Trích & chọn lọc từ vựng cần học từ các đề học viên đã làm (trả JSON)."""
    import json as _json
    lang = (payload.get("lang") or "vi").lower()
    # SỐ LƯỢNG: count = yêu cầu HV; min_count = sàn tối thiểu (mặc định 15). Trích nhắm tới max(count, min_count).
    min_count = int(payload.get("minCount", 15) or 15)
    count = int(payload.get("count", 0) or 0)
    target_count = max(count, min_count, 15)
    target = payload.get("target", "")
    source_limit = int(os.environ.get("VOCAB_SOURCE_CHARS") or (60000 if os.environ.get("OPENAI_API_KEY") else 18000))
    source = (payload.get("source", "") or "")[:max(8000, min(source_limit, 120000))]
    wrong = (payload.get("wrongContext", "") or "")[:1500]  # vùng liên quan câu sai
    # Từ đã có trong sổ — yêu cầu AI TRÁNH để mỗi lần ra từ MỚI (không cạn dần).
    _excl_in = payload.get("exclude") or []
    exclude = [str(x).strip().lower() for x in _excl_in if str(x).strip()][:200]

    # Loại mục HV muốn AI ưu tiên (mặc định cả 5). Lọc về danh sách hợp lệ.
    _ALL_KINDS = ["word", "phrasal_verb", "idiom", "collocation", "grammar"]
    _kinds_in = payload.get("kinds") or _ALL_KINDS
    kinds = [k for k in _ALL_KINDS if k in set(str(x).strip().lower().replace("-", "_") for x in _kinds_in)]
    if not kinds:
        kinds = _ALL_KINDS

    if not source.strip():
        return {"success": False, "error": "Chưa có dữ liệu đề để trích từ vựng."}

    # ĐỊNH NGHĨA: LUÔN tiếng Anh bất kể ngôn ngữ giao diện (yêu cầu sư phạm).
    meaning_key = "meaning_en"
    meaning_field = "short, clear English definition (learner-dictionary style)"
    _CAT_LABEL = {"word": "single words", "phrasal_verb": "phrasal verbs", "idiom": "idioms",
                  "collocation": "collocations / fixed phrases", "grammar": "grammar patterns"}
    _excl_clause = ""
    if exclude:
        _excl_clause = ("\nALREADY IN THE STUDENT'S NOTEBOOK — do NOT return any of these (return only NEW items):\n"
                        + ", ".join(exclude[:300]) + "\n")

    def _build_sys(kinds_list, n_target):
        all_set = (len(kinds_list) == 5)
        kinds_csv = ", ".join(kinds_list)
        if all_set:
            kinds_clause = (
                "You MUST return a BALANCED MIX across ALL FIVE categories (word, phrasal_verb, idiom, collocation, grammar) "
                "whenever the material allows. ACTIVELY HUNT for PHRASAL VERBS and COLLOCATIONS — do NOT return an all-single-words "
                "list. Target at least 2 phrasal verbs and 2 collocations if any genuinely appear in the material")
        else:
            kinds_clause = (
                f"Return ONLY items whose category is one of: {kinds_csv}. Do NOT return any item of any other category. "
                f"Spread items across just these chosen categories (aim for several of EACH)")
        return (
            "You are an IELTS language coach. Your ONLY source is the MATERIAL text the student actually studied (given below). "
            f"EXTRACT AT LEAST {n_target} (ideally {n_target + 6}) genuine language items worth memorising for IELTS. {kinds_clause}. "
            "READ THE WHOLE MATERIAL (it may contain several passages/tests) and mine ALL of it, not just the start.\n"
            "DO NOT EXTRACT (these are NOT lexical items, they are just a sentence's verb in some tense):\n"
            "- Bare verb conjugations / tense instances, e.g. 'have been grown', 'are generally cultivated', 'have made', "
            "'has been owned', 'was developed', 'is being built'. These have ZERO value as flashcards — SKIP them.\n"
            "- If a TENSE or STRUCTURE is genuinely worth learning, output it ONCE as a \"grammar\" item using an ABSTRACT "
            "PATTERN (e.g. 'have/has been + past participle (present perfect passive)') and quote the real sentence in source_sentence.\n"
            "A 'collocation' must be a REUSABLE lexical chunk (verb+noun like 'pose a threat', adj+noun, noun+preposition), "
            "NOT the main verb of a sentence in a particular tense.\n"
            "- AD-HOC COMBINATIONS: do NOT extract two words just because they sit next to each other. "
            "'carved deep', 'retreating glaciers', 'kettle ponds', 'low-lying areas', 'formative hothouse' are NOT vocabulary "
            "items. If such a phrase contains a valuable WORD, extract the SINGLE WORD instead (from 'carved deep' -> 'carve').\n"
            "- Only extract a MULTI-WORD item if it is an ESTABLISHED expression you could look up in a dictionary (a real phrasal "
            "verb, a real idiom, or a standard collocation/linking phrase).\n"
            "- Do NOT extract a word only meaningful as part of a fixed compound here (e.g. 'kettle' from 'kettle pond').\n"
            "PHRASAL VERBS — look hard: a base verb + a true ADVERBIAL PARTICLE (up/out/down/off/on/in/into/onto/away/back/over/through...) "
            "forming a dictionary headword ('carry out', 'set off', 'point out', 'put up with', 'break down', 'take on', 'bring about', "
            "'go into', 'look into', 'run into', 'turn into'). 'into'/'onto' ARE particles.\n"
            "PHRASAL VERB vs COLLOCATION — the #1 mistake, decide with THIS TEST:\n"
            "  (a) If the verb keeps its LITERAL meaning and the second word is a DEPENDENT PREPOSITION fixed to it, it is a "
            "PREPOSITIONAL VERB => category is \"collocation\", NOT phrasal_verb. Examples: 'depend on', 'rely on', 'focus on', "
            "'insist on', 'consist of', 'approve of', 'refer to', 'adhere to', 'contribute to', 'react to', 'cope with', "
            "'comply with', 'differ from', 'refrain from', 'account for', 'believe in', 'result in', 'specialise in', "
            "'embark on', 'delve into', 'preside over', 'worry about', 'guard against'. Here the particle CANNOT move after a "
            "noun object (you CANNOT say 'depend him on').\n"
            "  (b) It is a true PHRASAL VERB only if the particle is ADVERBIAL — it changes/completes the meaning idiomatically "
            "('give up' != give, 'carry out', 'put off', 'bring about') OR the particle can move after a noun object "
            "('carry out the plan' = 'carry the plan out'). Then category = phrasal_verb.\n"
            "IDIOM vs PHRASAL/COLLOCATION — classify as idiom only when the whole expression is figurative and dictionary-attested. "
            "Examples: 'keep an eye on', 'turn a blind eye to', 'break new ground', 'the tip of the iceberg'. "
            "Do NOT call normal verb+preposition chunks idioms.\n"
            "Verb + plain preposition ('refer to', 'belong to', 'look at') is NOT a phrasal verb — it is a collocation (or, if a "
            "basic A1/A2 verb like 'look at'/'listen to', drop it as worthless).\n"
            "GRAMMAR — ACTIVELY look for notable structures (inversion after a negative adverbial, passive reporting "
            "'be said/believed to + verb', cleft 'It was X that ...', correlative comparatives 'the more ... the more ...', "
            "conditionals, reduced relative/participle clauses). Output the ABSTRACT pattern as 'word' and quote the real sentence.\n"
            "Give a ROUGH \"category\" guess. A second specialist pass REFINES & VALIDATES it, so extract faithfully.\n"
            "ABSOLUTE RULES (anti-fabrication) — violating ANY makes the item INVALID:\n"
            "- The 'word' itself MUST be a verbatim substring of the MATERIAL (allow only normal inflection -s/-ed/-ing). "
            "The 'source_sentence' you quote MUST CONTAIN that exact 'word'.\n"
            "- Do NOT output a phrase because a SYNONYM appears (text 'pioneered' -> do NOT output 'break new ground'). "
            "If the full expression is not literally in the MATERIAL, it does not exist for you.\n"
            "- For grammar, the pattern MUST be evidenced by a real sentence in the MATERIAL.\n"
            "- NEVER invent, guess or pull from general knowledge.\n"
            "- Every item MUST include \"source_sentence\": a VERBATIM 6-25 word quote from the MATERIAL containing the exact 'word'.\n"
            "Prioritise items linked to questions the student got wrong, and high IELTS reuse value. Avoid trivial A1/A2 words and proper nouns.\n"
            "Respond as a JSON object {\"items\": [ ... ]} and nothing else. Each item MUST have keys: "
            "\"category\" (word|phrasal_verb|idiom|collocation|grammar), \"word\", \"source_sentence\", "
            "\"phonetic\" (IPA — may be empty for collocations/idioms/grammar), \"pos\", "
            f"\"{meaning_key}\" (the {meaning_field}. For a COLLOCATION, give a CONCRETE definition of the WHOLE chunk and how it is "
            "used — what it means + typical context/partner words (e.g. 'pose a threat = to be a danger to sb/sth; used with risks, "
            "species, security'); NEVER leave it vague, circular or empty. For grammar & idiom, explain the meaning AND when to use it), "
            "\"example\" (a natural English sentence using the item), \"cefr\" (B1/B2/C1...)."
        )

    _ext_model = os.environ.get("GROQ_EXTRACT_MODEL", "openai/gpt-oss-120b")
    _err_holder = [""]   # giữ lỗi Groq thật để báo cho HV (RATE_LIMIT, key sai, v.v.)

    def _extract(kinds_list, n_target, extra_exclude=None, classify=True):
        """Một lượt trích (giai đoạn A) -> classify -> normalize. Trả list item đã chuẩn hóa (chưa ground).
        None = lỗi Groq (xem _err_holder[0]); [] = chạy được nhưng rỗng."""
        sysp = _build_sys(kinds_list, n_target)
        dynamic_exclude = []
        for value in [*(exclude or []), *(extra_exclude or [])]:
            normalized = re.sub(r"\s+", " ", str(value or "").strip().lower())
            if normalized and normalized not in dynamic_exclude:
                dynamic_exclude.append(normalized)
        dynamic_clause = ""
        if dynamic_exclude:
            dynamic_clause = ("\nDO NOT REPEAT ANY ITEM ALREADY RETURNED OR ALREADY IN THE NOTEBOOK. "
                              "Treat the following normalized forms as unavailable:\n"
                              + ", ".join(dynamic_exclude[:500]) + "\n")
        userp = (
            f"Target band: {target}\n\nMATERIAL THE STUDENT STUDIED (your ONLY allowed source):\n\"\"\"\n{source}\n\"\"\"\n\n"
            f"AREAS LINKED TO MISTAKES:\n{wrong or '(none)'}\n{_excl_clause}{dynamic_clause}\n"
            f"Return the JSON object now with AT LEAST {n_target} items, all quotable verbatim from the MATERIAL above."
        )
        # A large chain of high-reasoning calls exceeds the 60-second serverless window.
        # One focused extraction still has enough room for every requested flashcard.
        max_tokens = max(4096, min(7000, int(n_target) * 220))
        txt, e = _vocab_chat(
            sysp, userp, max_tokens, json_mode=True, model=_ext_model,
            reasoning_effort=os.environ.get("VOCAB_EXTRACT_REASONING", "medium"),
        )
        if e:
            _err_holder[0] = e
            return None
        if not txt:
            return []
        arr = _parse_ai_items(txt)
        if not isinstance(arr, list):
            return []
        # Pass classify riêng tăng precision phrasal/collocation/idiom. Tắt bằng VOCAB_CLASSIFY_PASS=0 nếu cần tiết kiệm.
        if classify and os.environ.get("VOCAB_CLASSIFY_PASS", "1").strip().lower() not in ("0", "false", "off", "no"):
            arr = _classify_vocab_items(arr, lang)
        return _normalize_vocab_items(arr)

    # ===== PASS CHÍNH =====
    started_at = time.monotonic()
    items = _extract(kinds, target_count)
    if items is None:
        msg = _friendly_err(_err_holder[0] or "AI service error", lang)
        if _err_holder[0] == "RATE_LIMIT":
            # Chẩn đoán: server có NHẬN được key của provider nào không (giúp biết đã set env + redeploy chưa).
            provs = []
            gk = sum(1 for k in ["GROQ_API_KEY", "GROQ_API_KEY_2", "GROQ_API_KEY_3", "GROQ_API_KEY_4"] if os.environ.get(k))
            if gk: provs.append(f"groq×{gk}")
            if os.environ.get("CEREBRAS_API_KEY"): provs.append("cerebras")
            if os.environ.get("GEMINI_API_KEY") or os.environ.get("GOOGLE_API_KEY"): provs.append("gemini")
            if os.environ.get("OPENAI_API_KEY"): provs.insert(0, "openai")
            msg += f"  ·  providers loaded: {', '.join(provs) if provs else 'chưa thấy provider key — kiểm tra env + redeploy'}"
        return {"success": False, "error": msg}

    def _wkey(it):
        return re.sub(r"\s+", " ", re.sub(r"[^a-z0-9 ]", " ", str(it.get("word", "") or "").strip().lower())).strip()

    def _dedupe(items):
        unique = []
        seen = set()
        for item in items or []:
            if not isinstance(item, dict):
                continue
            key = _wkey(item)
            if not key or key in seen:
                continue
            seen.add(key)
            unique.append(item)
        return unique

    def _ground_unique(items):
        grounded, rejected = _ground_vocab_items(items, source)
        unique = _dedupe(grounded)
        return unique, rejected + max(0, len(grounded) - len(unique))

    items = _dedupe(items)

    # ===== PASS BÙ PHỦ LOẠI: đẩy mạnh phrasal/collocation/idiom (hay bị lép so với từ đơn) =====
    # Bù khi loại được yêu cầu mà còn DƯỚI ngưỡng — không chỉ khi = 0.
    from collections import Counter as _Counter
    _need = {"phrasal_verb": 2, "collocation": 2, "idiom": 1}
    _cnt = _Counter(it.get("category") for it in items if isinstance(it, dict))
    under = [k for k, m in _need.items() if k in kinds and _cnt.get(k, 0) < m]
    if under and os.environ.get("VOCAB_CATEGORY_REFILL", "0").strip().lower() in ("1", "true", "on", "yes"):
        extra = _extract(under, max(6, len(under) * 4)) or []
        seen = {_wkey(it) for it in items}
        for it in extra:
            if isinstance(it, dict) and it.get("category") in under and _wkey(it) and _wkey(it) not in seen:
                seen.add(_wkey(it)); items.append(it)

    # Nếu HV chọn nhóm con: tuyệt đối không trả lẫn nhóm khác. Thiếu thì báo shortfall trung thực.
    if len(kinds) < 5:
        _kept = [it for it in items if isinstance(it, dict) and it.get("category") in set(kinds)]
        items = _kept

    # ===== BỘ DÒ NGỮ PHÁP TẤT ĐỊNH (định nghĩa English) =====
    if "grammar" in kinds and source:
        _existing_keys = set()
        for it in items:
            if isinstance(it, dict) and it.get("category") == "grammar":
                _w = str(it.get("word", "")).lower()
                for _k, _rx, *_ in _GRAMMAR_PATTERNS:
                    if _rx.search(_w):
                        _existing_keys.add(_k)
        for g in _detect_grammar(source, "en"):
            if g.get("_grammar_key") not in _existing_keys:
                _existing_keys.add(g.pop("_grammar_key", None))
                items.append(g)

    # ===== CHỐNG BỊA + BÙ ĐỦ SỐ LƯỢNG =====
    # Grounding/classification có thể loại nhiều candidate. Gọi thêm vài lượt với danh sách
    # loại trừ động để không trả 1 item khi học viên yêu cầu 25, nhưng vẫn không bịa ngoài đề.
    verified, dropped = _ground_unique(items)
    items = verified
    max_refill_attempts = max(0, min(2, int(os.environ.get("VOCAB_MAX_REFILL_ATTEMPTS", "1") or 1)))
    for _attempt in range(max_refill_attempts):
        if len(items) >= target_count:
            break
        # Leave time for FastAPI/Vercel to serialize a proper response instead of hard timing out.
        if time.monotonic() - started_at > 28:
            break
        deficit = target_count - len(items)
        # Never let refill candidates bypass the specialist classifier. If that provider fails,
        # _normalize_vocab_items still applies the conservative deterministic fallback.
        refill = _extract(kinds, max(6, deficit + 6), extra_exclude=[_wkey(it) for it in items], classify=True)
        if not refill:
            break
        refill, refill_dropped = _ground_unique(refill)
        dropped += refill_dropped
        existing_keys = {_wkey(it) for it in items}
        additions = [it for it in refill if _wkey(it) not in existing_keys]
        if not additions:
            break
        items.extend(additions)

    # Refill cũng phải qua đúng bộ lọc nhóm; không để candidate bị phân loại lại lọt sang tab khác.
    if len(kinds) < 5:
        items = [it for it in items if isinstance(it, dict) and it.get("category") in set(kinds)]

    return {
        "success": True,
        "items": items[:target_count],
        "requested": target_count,
        "returned": min(len(items), target_count),
        "shortfall": max(0, target_count - len(items)),
        "dropped": dropped,
    }

def _normalize_audio_url(url: str) -> str:
    """Chuyển link Google Drive 'xem' thành link tải trực tiếp."""
    m = re.search(r"drive\.google\.com/file/d/([^/]+)", url)
    if m:
        return f"https://drive.google.com/uc?export=download&id={m.group(1)}"
    m = re.search(r"drive\.google\.com/open\?id=([^&]+)", url)
    if m:
        return f"https://drive.google.com/uc?export=download&id={m.group(1)}"
    return url


def _build_multipart(fields: dict, file_field: str, filename: str, file_bytes: bytes, file_mime: str):
    """Tạo body multipart/form-data thủ công cho upload audio."""
    import uuid as _uuid
    boundary = "----IELTSOS" + _uuid.uuid4().hex
    pre = []
    for k, v in fields.items():
        pre.append(f"--{boundary}\r\nContent-Disposition: form-data; name=\"{k}\"\r\n\r\n{v}\r\n")
    head = "".join(pre).encode("utf-8")
    file_head = (f"--{boundary}\r\nContent-Disposition: form-data; name=\"{file_field}\"; "
                 f"filename=\"{filename}\"\r\nContent-Type: {file_mime}\r\n\r\n").encode("utf-8")
    tail = f"\r\n--{boundary}--\r\n".encode("utf-8")
    body = head + file_head + file_bytes + tail
    return body, boundary


def _fmt_ts(sec: float) -> str:
    """Giây -> 'm:ss' hoặc 'h:mm:ss' (khớp format mốc transcript docx gốc)."""
    s = max(0, int(sec))
    h, m, ss = s // 3600, (s % 3600) // 60, s % 60
    return f"{h}:{m:02d}:{ss:02d}" if h else f"{m}:{ss:02d}"


def _segments_to_marked_transcript(segments, block_sec: float = 8.0) -> str:
    """Ghép segment Whisper (verbose_json) thành transcript CÓ MỐC THỜI GIAN TUYỆT ĐỐI:
    (m:ss - m:ss)\\ntext... — mốc lấy thẳng từ máy, không phải AI đoán."""
    blocks = []
    cur_start, cur_end, cur_texts = None, None, []
    for seg in segments or []:
        try:
            st, en = float(seg.get("start", 0)), float(seg.get("end", 0))
        except Exception:
            continue
        txt = str(seg.get("text", "") or "").strip()
        if not txt:
            continue
        if cur_start is None:
            cur_start, cur_end, cur_texts = st, en, [txt]
        elif (en - cur_start) > block_sec:
            blocks.append((cur_start, cur_end, " ".join(cur_texts)))
            cur_start, cur_end, cur_texts = st, en, [txt]
        else:
            cur_end = en
            cur_texts.append(txt)
    if cur_start is not None:
        blocks.append((cur_start, cur_end, " ".join(cur_texts)))
    return "\n\n".join(f"({_fmt_ts(a)} - {_fmt_ts(b)})\n{t}" for a, b, t in blocks)


_TS_MARKER_RE = re.compile(r"\((\d{1,2}):(\d{2})(?::(\d{2}))?\s*-\s*\d{1,2}:\d{2}(?::\d{2})?\)")
_TS_CITE_RE = re.compile(r"\s*\[(\d{1,2}):(\d{2})(?::(\d{2}))?\]")
_TS_RANGE_RE = re.compile(
    r"[\[(]\s*(\d{1,2}:\d{2}(?::\d{2})?)\s*(?:-|–|—|to)\s*\d{1,2}:\d{2}(?::\d{2})?\s*[\])]")
_TS_PAREN_CITE_RE = re.compile(r"\((\d{1,2}:\d{2}(?::\d{2})?)\)")
_TS_BLOCK_RE = re.compile(
    r"\((\d{1,2}):(\d{2})(?::(\d{2}))?\s*-\s*\d{1,2}:\d{2}(?::\d{2})?\)\s*\n?([\s\S]*?)(?=\n\s*\(\d{1,2}:\d{2}(?::\d{2})?\s*-|\Z)")


def _ts_to_sec(a: str, b: str, c) -> int:
    return (int(a) * 3600 + int(b) * 60 + int(c)) if c else (int(a) * 60 + int(b))


def _normalized_timestamp_text(value: str) -> str:
    return re.sub(r"[^a-z0-9]+", " ", str(value or "").lower()).strip()


def _answer_anchor_seconds(context: str, correct: str, answer_sequence=None, question_index=None):
    """Return only a transcript block that can be tied to the answer text and its listening order."""
    blocks = []
    for match in _TS_BLOCK_RE.finditer(context or ""):
        blocks.append((_ts_to_sec(match.group(1), match.group(2), match.group(3)), _normalized_timestamp_text(match.group(4))))
    if not blocks:
        return None

    candidates = {
        _normalized_timestamp_text(part)
        for part in re.split(r"\s*(?:/|;|\||\bor\b)\s*", str(correct or ""), flags=re.I)
    }
    candidates = {candidate for candidate in candidates if len(candidate) >= 3 and not candidate.isdigit()}
    for candidate in sorted(candidates, key=len, reverse=True):
        matches = [sec for sec, block in blocks if re.search(r"(?:^|\s)" + re.escape(candidate) + r"(?:\s|$)", block)]
        if len(matches) == 1:
            return matches[0]

    # When the same answer occurs more than once, use the ordered answer list from
    # the current listening part to select the first monotonic match for this item.
    # This prevents the model from choosing a later repeated occurrence.
    if isinstance(answer_sequence, list) and isinstance(question_index, int) and 0 <= question_index < len(answer_sequence):
        cursor = -1
        for idx, sequence_answer in enumerate(answer_sequence[:question_index + 1]):
            sequence_candidates = {
                _normalized_timestamp_text(part)
                for part in re.split(r"\s*(?:/|;|\||\bor\b)\s*", str(sequence_answer or ""), flags=re.I)
            }
            sequence_candidates = {candidate for candidate in sequence_candidates if len(candidate) >= 3 and not candidate.isdigit()}
            occurrences = [block_idx for block_idx, (_, block) in enumerate(blocks)
                           if any(re.search(r"(?:^|\s)" + re.escape(candidate) + r"(?:\s|$)", block)
                                  for candidate in sorted(sequence_candidates, key=len, reverse=True))]
            occurrences = [block_idx for block_idx in occurrences if block_idx > cursor]
            if not occurrences:
                continue
            cursor = occurrences[0]
            if idx == question_index:
                return blocks[cursor][0]
    return None


def _filter_fake_timestamps(answer: str, context: str, correct: str = "", lang: str = "vi", answer_sequence=None, question_index=None) -> str:
    """Keep only real seek markers and prefer a unique answer-bearing transcript block when available."""
    answer = _TS_RANGE_RE.sub(lambda m: "[" + m.group(1) + "]", answer or "")
    answer = _TS_PAREN_CITE_RE.sub(lambda m: "[" + m.group(1) + "]", answer)
    allowed = {_ts_to_sec(m.group(1), m.group(2), m.group(3)) for m in _TS_MARKER_RE.finditer(context or "")}
    if not allowed:
        return _TS_CITE_RE.sub("", answer)

    answer_anchor = _answer_anchor_seconds(context, correct, answer_sequence, question_index)
    if answer_anchor is not None:
        anchor = "[" + _fmt_ts(answer_anchor) + "]"
        if _TS_CITE_RE.search(answer):
            used_anchor = False
            def _replace_with_anchor(m):
                nonlocal used_anchor
                if used_anchor:
                    return ""
                used_anchor = True
                token = m.group(0)
                leading_space = token[:len(token) - len(token.lstrip())]
                return leading_space + anchor
            return _TS_CITE_RE.sub(_replace_with_anchor, answer)
        return answer.rstrip() + ("\n\nListen again: " if lang == "en" else "\n\nNghe lại: ") + anchor

    # A marker that is merely a valid transcript boundary is not evidence that it
    # contains this answer. A missing marker is safer than a confidently wrong one.
    return _TS_CITE_RE.sub("", answer)


@app.post("/api/ai_transcribe")
async def ai_transcribe(payload: Dict[str, Any] = Body(...)):
    """Chép lời audio bài Listening bằng Groq Whisper (1 lần, nhanh). Lưu vào đề để dùng lại."""
    import urllib.request as _urlreq
    import urllib.error as _urlerr

    lang = (payload.get("lang") or "vi").lower()
    audio_url = (payload.get("audioUrl") or "").strip()
    if not audio_url:
        return {"success": False, "error": "Đề này chưa có link audio."}

    api_key = os.environ.get("GROQ_API_KEY")
    if not api_key:
        return {"success": False, "error": "Server chưa cấu hình GROQ_API_KEY."}

    dl_url = _normalize_audio_url(audio_url)
    try:
        req = _urlreq.Request(dl_url, headers={"User-Agent": "Mozilla/5.0"})
        with _urlreq.urlopen(req, timeout=45) as r:
            raw = r.read()
            mime = (r.headers.get("Content-Type") or "audio/mpeg").split(";")[0].strip()
    except Exception as e:
        return {"success": False, "error": f"Không tải được audio từ link (đảm bảo link công khai): {e}"}

    if len(raw) > 25 * 1024 * 1024:
        return {"success": False, "error": f"File audio {len(raw)//(1024*1024)}MB > 25MB (giới hạn Whisper free). Hãy nén bitrate thấp hơn (vd 64kbps mono) hoặc tách section."}
    if "octet" in mime or not any(k in mime for k in ["audio", "mpeg", "ogg", "wav", "mp4", "m4a", "webm"]):
        mime = "audio/mpeg"

    model = os.environ.get("GROQ_WHISPER_MODEL", "whisper-large-v3-turbo")
    # verbose_json -> Whisper trả từng segment kèm start/end TUYỆT ĐỐI theo file audio
    # -> transcript nhúng mốc (m:ss - m:ss) chính xác máy, hết cảnh AI bịa timestamp.
    body, boundary = _build_multipart(
        {"model": model, "response_format": "verbose_json", "language": "en", "temperature": "0"},
        "file", "audio.mp3", raw, mime,
    )
    try:
        req = _urlreq.Request(
            "https://api.groq.com/openai/v1/audio/transcriptions", data=body, method="POST",
            headers={
                "Authorization": f"Bearer {api_key}",
                "Content-Type": f"multipart/form-data; boundary={boundary}",
                "User-Agent": "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/124.0 Safari/537.36",
                "Accept": "application/json",
            },
        )
        with _urlreq.urlopen(req, timeout=170) as r:
            _raw_resp = r.read().decode("utf-8").strip()
        import json as _json
        try:
            _vj = _json.loads(_raw_resp)
            segs = _vj.get("segments") or []
            text = _segments_to_marked_transcript(segs) if segs else str(_vj.get("text", "") or "").strip()
        except Exception:
            text = _raw_resp  # fallback: server trả text thường
    except _urlerr.HTTPError as e:
        detail = ""
        try: detail = e.read().decode("utf-8")[:300]
        except Exception: pass
        if e.code == 429:
            return {"success": False, "error": _friendly_err("RATE_LIMIT", lang)}
        return {"success": False, "error": f"Lỗi Groq {e.code}: {detail}"}
    except Exception as e:
        return {"success": False, "error": f"Lỗi chép lời: {e}"}

    if not text:
        return {"success": False, "error": "Không nghe được nội dung (audio rỗng?)."}
    return {"success": True, "transcript": text}


@app.get("/api/health")
async def health_check():
    return {"status": "ok", "timestamp": time.time()}
